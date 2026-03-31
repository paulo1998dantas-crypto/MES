import os
import pandas as pd
import io
import hashlib
import hmac
import random
import re
import secrets
import unicodedata
from pathlib import Path
from fastapi import FastAPI, Request, Depends, Body, UploadFile, File, Form
from fastapi.responses import StreamingResponse, RedirectResponse, HTMLResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from starlette.middleware.sessions import SessionMiddleware
from sqlalchemy.orm import Session
from sqlalchemy import or_, func, cast, String, text, inspect
import uvicorn
from zoneinfo import ZoneInfo
import datetime

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Configuração de diretórios e templates
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
templates = Jinja2Templates(directory=os.path.join(BASE_DIR, "templates"))
UPLOADS_DIR = Path(BASE_DIR) / "uploads"
OS_UPLOAD_DIR = UPLOADS_DIR / "ordens_servico"
OS_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

import database, models

LOCAL_TZ = ZoneInfo("America/Sao_Paulo")

# Inicialização do banco de dados
database.Base.metadata.create_all(bind=database.engine)

# Garante novas colunas sem migração formal
def ensure_columns():
    insp = inspect(database.engine)
    dialect = database.engine.dialect.name
    ts_tz = "TIMESTAMP WITH TIME ZONE" if dialect == "postgresql" else "DATETIME"
    bool_true = "BOOLEAN DEFAULT TRUE" if dialect == "postgresql" else "BOOLEAN DEFAULT 1"

    def column_names(table):
        if not insp.has_table(table):
            return set()
        return {c["name"] for c in insp.get_columns(table)}

    def add_column_if_missing(conn, table, column, definition):
        if column in column_names(table):
            return
        conn.execute(text(f"ALTER TABLE {table} ADD COLUMN {column} {definition}"))

    with database.engine.begin() as conn:
        add_column_if_missing(conn, "veiculos", "ar_condicionado", "VARCHAR")
        add_column_if_missing(conn, "veiculos", "ordem", "INTEGER")
        add_column_if_missing(conn, "veiculos", "ativo", bool_true)
        add_column_if_missing(conn, "veiculos", "cj_bco", "VARCHAR")
        add_column_if_missing(conn, "veiculos", "cliente", "VARCHAR")
        add_column_if_missing(conn, "veiculos", "destino", "VARCHAR")
        add_column_if_missing(conn, "veiculos", "localizacao", "VARCHAR")
        add_column_if_missing(conn, "veiculos", "banco_presente", "VARCHAR")
        add_column_if_missing(conn, "veiculos", "banco_comentario", "VARCHAR")
        conn.execute(text("UPDATE veiculos SET ativo = 1 WHERE ativo IS NULL"))

        add_column_if_missing(conn, "apontamentos", "responsavel", "VARCHAR")
        add_column_if_missing(conn, "apontamentos", "inicio", ts_tz)
        add_column_if_missing(conn, "apontamentos", "termino", ts_tz)
        add_column_if_missing(conn, "apontamentos", "localizacao", "VARCHAR")
        add_column_if_missing(conn, "apontamentos", "observacao", "VARCHAR")

        add_column_if_missing(conn, "bom_itens", "qtd_apontada", "FLOAT")

        add_column_if_missing(conn, "empenhos", "qtd_empenhada", "FLOAT")

        hist_cols = column_names("historico")
        add_column_if_missing(conn, "historico", "responsavel", "VARCHAR")
        add_column_if_missing(conn, "historico", "inicio", ts_tz)
        add_column_if_missing(conn, "historico", "termino", ts_tz)
        add_column_if_missing(conn, "historico", "localizacao", "VARCHAR")
        add_column_if_missing(conn, "historico", "observacao", "VARCHAR")
        if "data_apontamento" in hist_cols and dialect == "postgresql":
            conn.execute(text("ALTER TABLE historico ALTER COLUMN data_apontamento TYPE TIMESTAMP WITH TIME ZONE"))

ensure_columns()

app = FastAPI()
app.add_middleware(
    SessionMiddleware,
    secret_key=os.getenv("SESSION_SECRET", "local-dev-session-secret"),
    same_site="lax",
)
app.mount("/static", StaticFiles(directory=os.path.join(BASE_DIR, "static")), name="static")

PERFIS_USUARIO = [
    "ADM",
    "LIDER",
    "VIDROS",
    "REVESTIMENTO",
    "DESMONTAGEM",
    "ELETRICA",
    "BANCO",
    "PREPARACAO",
    "EXPEDICAO",
    "SERRALHERIA",
    "LIBERACAO",
]

DEFAULT_ADMIN_LOGIN = os.getenv("DEFAULT_ADMIN_LOGIN", "admin")
DEFAULT_ADMIN_PASSWORD = os.getenv("DEFAULT_ADMIN_PASSWORD", "admin123")

MANAGEMENT_PROFILES = {"ADM", "LIDER"}

POSTOS_TRABALHO = {
    "CORTE_VIDRO": {"label": "Posto Corte de Vidro", "etapa": "VIDROS", "modo": "operacao"},
    "DESMONTAGEM": {"label": "Posto Desmontagem", "etapa": "DESMONT", "modo": "operacao"},
    "REVESTIMENTO_1": {"label": "Posto Revestimento 1", "etapa": "REVEST", "modo": "operacao"},
    "REVESTIMENTO_2": {"label": "Posto Revestimento 2", "etapa": "REVEST", "modo": "operacao"},
    "REVESTIMENTO_3": {"label": "Posto Revestimento 3", "etapa": "REVEST", "modo": "operacao"},
    "REVESTIMENTO_4": {"label": "Posto Revestimento 4", "etapa": "REVEST", "modo": "operacao"},
    "TERCERIZACAO": {"label": "Posto Tercerização", "etapa": "SERRA.", "modo": "operacao"},
    "SERRALHERIA_BANCOS": {"label": "Posto Serralheria Bancos", "etapa": "SERRA.", "modo": "operacao"},
    "ELETRICA": {"label": "Posto Elétrica", "etapa": "ELÉTRICA", "modo": "operacao"},
    "BANCOS": {"label": "Posto Montagem Bancos", "etapa": "BCO", "modo": "operacao"},
    "PREPARACAO": {"label": "Posto Preparação", "etapa": "PREP", "modo": "checklist", "bom_tipo": "PREPARACAO"},
    "EXPEDICAO": {"label": "Posto Expedição", "etapa": "EXPE.", "modo": "checklist", "bom_tipo": "EXPEDICAO", "permite_empenho": True},
    "LIBERACAO": {"label": "Posto Liberação", "etapa": "LIBERA.", "modo": "operacao"},
}

SELECTION_POSTS_BY_PROFILE = {
    "VIDROS": ["CORTE_VIDRO"],
    "DESMONTAGEM": ["DESMONTAGEM"],
    "ELETRICA": ["ELETRICA"],
    "REVESTIMENTO": ["REVESTIMENTO_1", "REVESTIMENTO_2", "REVESTIMENTO_3", "REVESTIMENTO_4"],
    "BANCO": ["BANCOS"],
    "SERRALHERIA": ["TERCERIZACAO", "SERRALHERIA_BANCOS"],
}

DIRECT_POST_BY_PROFILE = {
    "PREPARACAO": "PREPARACAO",
    "EXPEDICAO": "EXPEDICAO",
    "LIBERACAO": "LIBERACAO",
}

BOM_TIPOS = {
    "PREPARACAO": "B.O.M. Preparação",
    "EXPEDICAO": "B.O.M. Expedição",
}

ETAPAS_GESTAO_AUTOMATICA = {
    "PREP": "Status controlado automaticamente pelo checklist da B.O.M. da Preparação.",
    "EXPE.": "Status controlado automaticamente pelo checklist da B.O.M. da Expedição.",
}

ETAPAS_PRODUCAO = [
    "VIDROS",
    "A/C",
    "PREP",
    "SERRA.",
    "EXPE.",
    "DESMONT",
    "ELÉTRICA",
    "REVEST",
    "BCO",
    "ACESSÓ.",
    "PLOTA.",
    "LIBERA."
]

ETAPAS_STATUS_ATUAL = ["VIDROS", "A/C", "DESMONT", "REVEST", "BCO", "LIBERA."]

ETAPAS_FILTRO = [e for e in ETAPAS_PRODUCAO if e != "A/C"] + ["GE", "CLIM"]

LOCALIZACOES = [
    "Pátio",
    "J I",
    "Linha",
    "Tenda",
    "R1",
    "R2",
    "R3",
    "R4",
    "R5",
    "R6",
    "R7",
    "R8",
    "R9",
    "R10",
    "R11",
]

def parse_local_dt(value):
    if value is None or (hasattr(pd, "isna") and pd.isna(value)) or value == "":
        return None
    if isinstance(value, pd.Timestamp):
        value = value.to_pydatetime()
    if isinstance(value, datetime.datetime):
        dt = value
    else:
        try:
            dt = datetime.datetime.fromisoformat(str(value))
        except ValueError:
            return None
    if dt.tzinfo is None:
        return dt.replace(tzinfo=LOCAL_TZ)
    return dt.astimezone(LOCAL_TZ)

def to_excel_dt(value):
    if not value:
        return None
    if value.tzinfo is None:
        return value
    return value.astimezone(LOCAL_TZ).replace(tzinfo=None)

def to_input_dt(value):
    if not value:
        return ""
    if value.tzinfo is not None:
        value = value.astimezone(LOCAL_TZ)
    return value.strftime("%Y-%m-%dT%H:%M")

def normalize_login(value: str) -> str:
    return str(value or "").strip().lower()


def normalize_profile(value: str) -> str:
    perfil = str(value or "").strip().upper()
    return perfil if perfil in PERFIS_USUARIO else ""


def hash_password(password: str, salt: str = None) -> str:
    salt = salt or secrets.token_hex(16)
    pwd_hash = hashlib.pbkdf2_hmac(
        "sha256",
        str(password).encode("utf-8"),
        salt.encode("utf-8"),
        390000,
    )
    return f"{salt}${pwd_hash.hex()}"


def verify_password(password: str, stored_hash: str) -> bool:
    if not stored_hash or "$" not in stored_hash:
        return False
    salt, saved = stored_hash.split("$", 1)
    candidate = hash_password(password, salt)
    return hmac.compare_digest(candidate, f"{salt}${saved}")


def build_session_user(user: models.Usuario) -> dict:
    return {
        "id": user.id,
        "nome": user.nome,
        "login": user.login,
        "perfil": user.perfil,
    }


def get_current_user(request: Request):
    user = request.session.get("user") or {}
    if not isinstance(user, dict):
        return None
    if not user.get("id"):
        return None
    return user


def get_user_name(request: Request):
    user = get_current_user(request)
    return (user or {}).get("nome", "")


def get_user_profile(request: Request):
    user = get_current_user(request)
    return (user or {}).get("perfil", "")

def require_login(request: Request):
    return get_current_user(request)


def require_admin(request: Request):
    user = get_current_user(request)
    if not user:
        return None
    return user if user.get("perfil") == "ADM" else None


def can_export_expedicao_lancamentos(request: Request) -> bool:
    perfil = normalize_profile(get_user_profile(request))
    return perfil in {"ADM", "EXPEDICAO"}


def active_vehicle_filter():
    return models.Veiculo.ativo.is_(True)


def active_vehicle_query(db: Session):
    return db.query(models.Veiculo).filter(active_vehicle_filter())


def get_vehicle_by_chassi(db: Session, chassi: str, active_only: bool = True):
    query = db.query(models.Veiculo).filter(
        func.trim(cast(models.Veiculo.chassi, String)) == str(chassi).strip()
    )
    if active_only:
        query = query.filter(active_vehicle_filter())
    return query.first()


def get_active_chassis(db: Session):
    return [
        str(chassi).strip()
        for (chassi,) in active_vehicle_query(db).with_entities(models.Veiculo.chassi).all()
    ]


def get_next_vehicle_order(db: Session) -> int:
    maior_ordem = active_vehicle_query(db).with_entities(func.max(models.Veiculo.ordem)).scalar()
    return int(maior_ordem or 0) + 1


def is_management_profile(perfil: str) -> bool:
    return str(perfil or "").upper() in MANAGEMENT_PROFILES


def is_management_user(request: Request) -> bool:
    return is_management_profile(get_user_profile(request))


def get_allowed_posts_for_profile(perfil: str):
    perfil_limpo = normalize_profile(perfil)
    if perfil_limpo in DIRECT_POST_BY_PROFILE:
        return [DIRECT_POST_BY_PROFILE[perfil_limpo]]
    return SELECTION_POSTS_BY_PROFILE.get(perfil_limpo, [])


def get_operator_home_url_for_profile(perfil: str) -> str:
    perfil_limpo = normalize_profile(perfil)
    if is_management_profile(perfil_limpo):
        return "/"
    if perfil_limpo in DIRECT_POST_BY_PROFILE:
        return f"/postos/{DIRECT_POST_BY_PROFILE[perfil_limpo]}"
    return "/postos"


def get_operator_home_url(request: Request) -> str:
    return get_operator_home_url_for_profile(get_user_profile(request))


def get_posto_config(posto: str):
    return POSTOS_TRABALHO.get(str(posto or "").strip().upper())


def get_management_locked_stages_for_profile(perfil: str) -> dict:
    perfil_limpo = normalize_profile(perfil)
    bloqueadas = {}
    if perfil_limpo == "LIDER":
        bloqueadas["PREP"] = ETAPAS_GESTAO_AUTOMATICA["PREP"]
        bloqueadas["EXPE."] = ETAPAS_GESTAO_AUTOMATICA["EXPE."]
    return bloqueadas


def get_posto_mode(posto: str) -> str:
    cfg = get_posto_config(posto)
    return (cfg or {}).get("modo", "operacao")


def can_access_posto(request: Request, posto: str) -> bool:
    if is_management_user(request):
        return True
    allowed = get_allowed_posts_for_profile(get_user_profile(request))
    return str(posto or "").strip().upper() in allowed


def can_access_chassi(request: Request, db: Session, chassi: str) -> bool:
    if is_management_user(request):
        return True
    allowed = get_allowed_posts_for_profile(get_user_profile(request))
    if not allowed:
        return False
    found = db.query(models.PostoSequencia).filter(
        models.PostoSequencia.posto.in_(allowed),
        func.trim(cast(models.PostoSequencia.chassi, String)) == str(chassi).strip(),
    ).first()
    return bool(found)


def ensure_default_admin():
    db = database.SessionLocal()
    try:
        if db.query(models.Usuario).count() > 0:
            return
        db.add(
            models.Usuario(
                nome="Administrador",
                login=DEFAULT_ADMIN_LOGIN,
                senha_hash=hash_password(DEFAULT_ADMIN_PASSWORD),
                perfil="ADM",
            )
        )
        db.commit()
    finally:
        db.close()


ensure_default_admin()

def render_login_page(request: Request, db: Session, erro: str = "", login_value: str = ""):
    return templates.TemplateResponse(
        request,
        "login.html",
        {
            "request": request,
            "erro": erro,
            "login_value": login_value,
        },
        status_code=400 if erro else 200,
    )


def render_user_management(
    request: Request,
    db: Session,
    erro: str = "",
    sucesso: str = "",
    form_data: dict = None,
):
    usuarios = db.query(models.Usuario).order_by(models.Usuario.nome.asc()).all()
    return templates.TemplateResponse(
        request,
        "usuarios.html",
        {
            "request": request,
            "current_user": require_login(request),
            "usuarios": usuarios,
            "perfis": PERFIS_USUARIO,
            "erro": erro,
            "sucesso": sucesso,
            "form_data": form_data or {},
        },
        status_code=400 if erro else 200,
    )


def get_apontamento_for_stage(db: Session, chassi: str, etapa: str):
    aponts = db.query(models.Apontamento).filter(
        func.trim(cast(models.Apontamento.chassi, String)) == str(chassi).strip()
    ).all()
    etapa_norm = normalize_etapa(etapa)
    for apont in aponts:
        if normalize_etapa(apont.etapa) == etapa_norm:
            return apont
    return None


def get_or_create_apontamento(db: Session, chassi: str, etapa: str):
    apont = get_apontamento_for_stage(db, chassi, etapa)
    if apont:
        return apont
    apont = models.Apontamento(
        chassi=str(chassi).strip(),
        etapa=etapa,
        status="NAO",
    )
    db.add(apont)
    db.flush()
    return apont


def build_operacao_status(apont):
    if not apont:
        return "AGUARDANDO"
    status = str(apont.status or "").strip().upper()
    if status in ["SIM", "S", "OK"]:
        return "FINALIZADO"
    if apont.inicio and apont.termino:
        return "PARADO"
    if apont.inicio:
        return "EM ANDAMENTO"
    return "AGUARDANDO"


def get_posto_cards(db: Session, posto: str):
    posto_key = str(posto or "").strip().upper()
    posto_cfg = get_posto_config(posto_key)
    if not posto_cfg:
        return []

    sequencias = db.query(models.PostoSequencia).filter(
        models.PostoSequencia.posto == posto_key
    ).order_by(models.PostoSequencia.sequencia.asc(), models.PostoSequencia.id.asc()).all()

    cards = []
    for item in sequencias:
        veiculo = get_vehicle_by_chassi(db, item.chassi)
        if not veiculo:
            continue
        apont = get_apontamento_for_stage(db, veiculo.chassi, posto_cfg["etapa"])
        ordem_servico = get_ordem_servico(db, veiculo.chassi)
        cards.append(
            {
                "sequencia": item.sequencia,
                "posto": posto_key,
                "posto_label": posto_cfg["label"],
                "etapa": posto_cfg["etapa"],
                "veiculo": veiculo,
                "apontamento": apont,
                "status_operacao": build_operacao_status(apont),
                "tem_ordem_servico": bool(ordem_servico),
            }
        )
    return cards


def resequence_posto(db: Session, posto: str, chassi: str, sequencia: int):
    posto_key = str(posto or "").strip().upper()
    chassi_key = str(chassi or "").strip()
    items = db.query(models.PostoSequencia).filter(
        models.PostoSequencia.posto == posto_key
    ).order_by(models.PostoSequencia.sequencia.asc(), models.PostoSequencia.id.asc()).all()

    current = None
    others = []
    for item in items:
        if str(item.chassi).strip() == chassi_key:
            current = item
        else:
            others.append(item)

    if current is None:
        current = models.PostoSequencia(posto=posto_key, chassi=chassi_key, sequencia=sequencia)
        db.add(current)

    insert_at = max(0, min(int(sequencia) - 1, len(others)))
    others.insert(insert_at, current)
    for idx, item in enumerate(others, start=1):
        item.sequencia = idx


def remove_posto_sequencia(db: Session, posto: str, chassi: str) -> bool:
    posto_key = str(posto or "").strip().upper()
    chassi_key = str(chassi or "").strip()
    item = db.query(models.PostoSequencia).filter(
        models.PostoSequencia.posto == posto_key,
        func.trim(cast(models.PostoSequencia.chassi, String)) == chassi_key,
    ).first()
    if not item:
        return False

    db.delete(item)
    db.flush()

    restantes = db.query(models.PostoSequencia).filter(
        models.PostoSequencia.posto == posto_key
    ).order_by(models.PostoSequencia.sequencia.asc(), models.PostoSequencia.id.asc()).all()

    for idx, restante in enumerate(restantes, start=1):
        restante.sequencia = idx
    return True


def render_sequenciamento_page(
    request: Request,
    db: Session,
    erro: str = "",
    sucesso: str = "",
    form_data: dict = None,
):
    grupos = []
    for posto_key, posto_cfg in POSTOS_TRABALHO.items():
        cards = get_posto_cards(db, posto_key)
        status_counts = {
            "AGUARDANDO": 0,
            "EM ANDAMENTO": 0,
            "PARADO": 0,
            "FINALIZADO": 0,
        }
        for card in cards:
            status = str(card.get("status_operacao", "")).strip().upper()
            status_counts.setdefault(status, 0)
            status_counts[status] += 1

        grupos.append(
            {
                "key": posto_key,
                "label": posto_cfg["label"],
                "cards": cards,
                "cards_count": len(cards),
                "modo": posto_cfg.get("modo", "operacao"),
                "modo_label": "Checklist B.O.M." if posto_cfg.get("modo") == "checklist" else "Operacao",
                "status_counts": status_counts,
            }
        )

    total_cards = sum(grupo["cards_count"] for grupo in grupos)
    postos_com_fila = sum(1 for grupo in grupos if grupo["cards_count"] > 0)
    postos_vazios = len(grupos) - postos_com_fila

    return templates.TemplateResponse(
        request,
        "sequenciamento.html",
        {
            "request": request,
            "current_user": require_login(request),
            "postos": POSTOS_TRABALHO,
            "bom_tipos": BOM_TIPOS,
            "grupos": grupos,
            "erro": erro,
            "sucesso": sucesso,
            "form_data": form_data or {},
            "resumo": {
                "total_postos": len(grupos),
                "postos_com_fila": postos_com_fila,
                "postos_vazios": postos_vazios,
                "total_cards": total_cards,
            },
        },
        status_code=400 if erro else 200,
    )


def get_posto_by_bom_tipo(tipo: str):
    tipo_limpo = str(tipo or "").strip().upper()
    for posto_key, posto_cfg in POSTOS_TRABALHO.items():
        if str(posto_cfg.get("bom_tipo", "")).upper() == tipo_limpo:
            return posto_key, posto_cfg
    return None, None


def registrar_historico_evento(db: Session, veiculo, etapa: str, status: str, responsavel: str, inicio=None, termino=None, observacao=""):
    db.add(
        models.Historico(
            chassi=veiculo.chassi if veiculo else "",
            modelo=veiculo.modelo if veiculo else "N/A",
            etapa=etapa,
            status=status,
            responsavel=responsavel,
            inicio=inicio,
            termino=termino,
            localizacao=veiculo.localizacao if veiculo else None,
            observacao=observacao,
        )
    )


def get_ordem_servico(db: Session, chassi: str):
    return db.query(models.OrdemServico).filter(
        func.trim(cast(models.OrdemServico.chassi, String)) == str(chassi).strip()
    ).first()


def remove_ordem_servico_arquivo(ordem_servico):
    if not ordem_servico:
        return
    caminho = Path(ordem_servico.caminho_arquivo)
    if caminho.exists():
        caminho.unlink()


def extract_docx_preview(caminho_arquivo: str):
    if not caminho_arquivo or not Path(caminho_arquivo).exists():
        return {"paragraphs": [], "tables": [], "available": False}
    if DocxDocument is None:
        return {"paragraphs": [], "tables": [], "available": False}

    doc = DocxDocument(caminho_arquivo)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
    return {"paragraphs": paragraphs, "tables": tables, "available": True}


def extract_docx_composition_items(caminho_arquivo: str):
    if not caminho_arquivo or not Path(caminho_arquivo).exists() or DocxDocument is None:
        return []

    doc = DocxDocument(caminho_arquivo)
    composition_table = None
    for table in doc.tables:
        if not table.rows:
            continue
        header_text = " ".join(safe_str(cell.text) for cell in table.rows[0].cells).upper()
        if "COMPOSICAO" in normalize_lookup_key(header_text):
            composition_table = table
            break

    if not composition_table or len(composition_table.rows) < 3:
        return []

    items = []
    for row in composition_table.rows[2:]:
        cells = [safe_str(cell.text) for cell in row.cells]
        if len(cells) < 4:
            continue
        cod_item, item_nome, qtd, unidade = cells[:4]
        if not cod_item and not item_nome:
            continue
        items.append(
            {
                "cod_item": cod_item,
                "item": item_nome,
                "descricao": f"Un.: {unidade}" if unidade else "",
                "qtd": qtd,
            }
        )
    return items


def build_bom_match_key(cod_item: str, item: str, descricao: str):
    codigo_key = normalize_lookup_key(cod_item)
    if codigo_key:
        return ("CODIGO", codigo_key)
    return (
        "TEXTO",
        normalize_lookup_key(item),
        normalize_lookup_key(descricao),
    )


def sync_bom_items_for_chassi(db: Session, tipo: str, chassi: str, rows):
    tipo_key = str(tipo or "").strip().upper()
    chassi_key = str(chassi or "").strip()
    rows = rows or []

    existing_items = get_bom_items(db, tipo_key, chassi_key)
    existing_by_key = {}
    for existing in existing_items:
        key = build_bom_match_key(existing.cod_item, existing.item, existing.descricao)
        existing_by_key.setdefault(key, []).append(existing)

    kept_ids = set()
    for row in rows:
        key = build_bom_match_key(row.get("cod_item", ""), row.get("item", ""), row.get("descricao", ""))
        current = None
        candidates = existing_by_key.get(key) or []
        if candidates:
            current = candidates.pop(0)

        if not current:
            current = models.BomItem(
                tipo=tipo_key,
                chassi=chassi_key,
                status="NAO",
            )
            db.add(current)
            db.flush()

        current.cod_item = row.get("cod_item", "")
        current.item = row.get("item", "")
        current.descricao = row.get("descricao", "")
        current.qtd = row.get("qtd", "")
        kept_ids.add(current.id)

    obsolete_ids = [item.id for item in existing_items if item.id not in kept_ids]
    if obsolete_ids:
        db.query(models.Empenho).filter(
            models.Empenho.bom_item_id.in_(obsolete_ids)
        ).delete(synchronize_session=False)
        db.query(models.BomItem).filter(
            models.BomItem.id.in_(obsolete_ids)
        ).delete(synchronize_session=False)

    return len(rows)


def get_bom_items(db: Session, tipo: str, chassi: str):
    return db.query(models.BomItem).filter(
        models.BomItem.tipo == str(tipo).strip().upper(),
        func.trim(cast(models.BomItem.chassi, String)) == str(chassi).strip(),
    ).order_by(models.BomItem.id.asc()).all()


def get_bom_item(db: Session, item_id: int):
    return db.query(models.BomItem).filter(models.BomItem.id == int(item_id)).first()


def sync_stage_from_bom(db: Session, tipo: str, chassi: str, responsavel: str):
    posto_key, posto_cfg = get_posto_by_bom_tipo(tipo)
    if not posto_cfg:
        return None

    veiculo = get_vehicle_by_chassi(db, chassi)
    if not veiculo:
        return None

    itens = get_bom_items(db, tipo, chassi)
    if not itens:
        return None

    apont = get_or_create_apontamento(db, chassi, posto_cfg["etapa"])
    status_anterior = str(apont.status or "").strip().upper()
    agora = datetime.datetime.now(LOCAL_TZ)
    todos_ok = all(normalize_status_value(item.status) in ["SIM", "N/A"] for item in itens)

    if not apont.inicio:
        apont.inicio = agora
    apont.responsavel = responsavel
    apont.localizacao = veiculo.localizacao

    if todos_ok:
        apont.status = "SIM"
        apont.termino = agora
    else:
        apont.status = "NAO"
        apont.termino = None

    if status_anterior != str(apont.status).strip().upper():
        registrar_historico_evento(
            db,
            veiculo,
            posto_cfg["etapa"],
            apont.status,
            responsavel,
            inicio=apont.inicio,
            termino=apont.termino,
            observacao="Atualizacao automatica via checklist B.O.M.",
        )
    return apont


def get_sequence_number(db: Session, posto: str, chassi: str):
    item = db.query(models.PostoSequencia).filter(
        models.PostoSequencia.posto == str(posto).strip().upper(),
        func.trim(cast(models.PostoSequencia.chassi, String)) == str(chassi).strip(),
    ).first()
    return item.sequencia if item else None


def dataframe_to_excel_response(df: pd.DataFrame, filename: str):
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        df.to_excel(writer, index=False)
    out.seek(0)
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def safe_str(value):
    if value is None or (hasattr(pd, "isna") and pd.isna(value)):
        return ""
    return str(value).strip()


def parse_quantity_value(value):
    raw = safe_str(value)
    if not raw:
        return None
    normalized = raw.replace(" ", "")
    if "," in normalized and "." in normalized:
        if normalized.rfind(",") > normalized.rfind("."):
            normalized = normalized.replace(".", "").replace(",", ".")
        else:
            normalized = normalized.replace(",", "")
    else:
        normalized = normalized.replace(",", ".")
    try:
        return float(normalized)
    except ValueError:
        return None


def format_quantity_value(value):
    if value is None:
        return "-"
    if abs(value - round(value)) < 1e-9:
        return str(int(round(value)))
    return f"{value:.2f}".rstrip("0").rstrip(".").replace(".", ",")


def format_quantity_input(value):
    if value is None:
        return ""
    if abs(value - round(value)) < 1e-9:
        return str(int(round(value)))
    return f"{value:.2f}".rstrip("0").rstrip(".")


def resolve_bom_item_pointed_quantity(item, total_lancado=0.0):
    if item is None:
        return float(total_lancado or 0.0)
    if item.qtd_apontada is not None:
        return float(item.qtd_apontada or 0.0)
    return float(total_lancado or 0.0)


def compute_consumption_status(qtd_prevista, qtd_consumida):
    if qtd_prevista is None:
        return "SEM_QTD_PREVISTA"
    diff = (qtd_consumida or 0.0) - qtd_prevista
    if diff > 1e-9:
        return "EXCEDENTE"
    if diff < -1e-9:
        return "FALTANTE"
    return "CONFORME"


def get_empenhos_for_items(db: Session, bom_items):
    item_ids = [item.id for item in bom_items]
    if not item_ids:
        return {}
    empenhos = db.query(models.Empenho).filter(
        models.Empenho.bom_item_id.in_(item_ids)
    ).order_by(models.Empenho.criado_em.asc(), models.Empenho.id.asc()).all()
    grouped = {}
    for empenho in empenhos:
        grouped.setdefault(empenho.bom_item_id, []).append(empenho)
    return grouped


def build_bom_item_empenho_summary(bom_items, empenhos_por_item):
    summary = {}
    for item in bom_items:
        empenhos_item = empenhos_por_item.get(item.id, [])
        for empenho in empenhos_item:
            empenho.qtd_empenhada_fmt = format_quantity_value(float(empenho.qtd_empenhada or 0))
            empenho.criado_em_fmt = (
                empenho.criado_em.astimezone(LOCAL_TZ).strftime("%d/%m/%Y %H:%M")
                if empenho.criado_em
                else "-"
            )
        total_lancado = sum(float(emp.qtd_empenhada or 0) for emp in empenhos_item)
        qtd_apontada = resolve_bom_item_pointed_quantity(item, total_lancado)
        qtd_prevista = parse_quantity_value(item.qtd)
        saldo = None if qtd_prevista is None else qtd_prevista - qtd_apontada
        summary[item.id] = {
            "empenhos": empenhos_item,
            "total_consumido": total_lancado,
            "total_consumido_fmt": format_quantity_value(total_lancado),
            "total_lancado": total_lancado,
            "total_lancado_fmt": format_quantity_value(total_lancado),
            "qtd_prevista": qtd_prevista,
            "qtd_prevista_fmt": format_quantity_value(qtd_prevista),
            "qtd_apontada": qtd_apontada,
            "qtd_apontada_fmt": format_quantity_value(qtd_apontada),
            "qtd_apontada_input": format_quantity_input(item.qtd_apontada),
            "saldo": saldo,
            "saldo_fmt": format_quantity_value(saldo),
            "status_consumo": compute_consumption_status(qtd_prevista, qtd_apontada),
        }
    return summary


def build_expedicao_export_rows(db: Session):
    active_chassis = get_active_chassis(db)
    if not active_chassis:
        return []

    itens = db.query(models.BomItem).filter(
        models.BomItem.tipo == "EXPEDICAO",
        models.BomItem.chassi.in_(active_chassis),
    ).order_by(models.BomItem.chassi.asc(), models.BomItem.id.asc()).all()
    itens_por_id = {item.id: item for item in itens}
    ordens_por_chassi = {
        ordem.chassi: ordem
        for ordem in db.query(models.OrdemServico).all()
    }

    grouped = {}
    for item in itens:
        ordem_servico = ordens_por_chassi.get(item.chassi)
        key = (
            safe_str(item.chassi),
            safe_str(item.cod_item),
            safe_str(item.item),
            safe_str(item.descricao),
        )
        row = grouped.setdefault(
            key,
            {
                "CHASSI": safe_str(item.chassi),
                "ORDEM_SERVICO_ARQUIVO": safe_str(ordem_servico.nome_arquivo if ordem_servico else ""),
                "COD_ITEM": safe_str(item.cod_item),
                "ITEM": safe_str(item.item),
                "DESCRICAO": safe_str(item.descricao),
                "QTD_PREVISTA_TOTAL": 0.0,
                "QTD_PREVISTA_ORIGINAL": [],
                "STATUS_CHECKLIST": "SIM",
                "RESPONSAVEL_CHECKLIST": [],
                "ULTIMA_ATUALIZACAO_CHECKLIST": None,
                "ITEM_IDS": [],
            },
        )
        qtd_prevista = parse_quantity_value(item.qtd)
        if qtd_prevista is not None:
            row["QTD_PREVISTA_TOTAL"] += qtd_prevista
        if safe_str(item.qtd):
            row["QTD_PREVISTA_ORIGINAL"].append(safe_str(item.qtd))
        row["ITEM_IDS"].append(item.id)
        if normalize_status_value(item.status) not in ["SIM", "N/A"]:
            row["STATUS_CHECKLIST"] = "PENDENTE"
        if safe_str(item.responsavel):
            row["RESPONSAVEL_CHECKLIST"].append(safe_str(item.responsavel))
        if item.atualizado_em and (
            row["ULTIMA_ATUALIZACAO_CHECKLIST"] is None
            or item.atualizado_em > row["ULTIMA_ATUALIZACAO_CHECKLIST"]
        ):
            row["ULTIMA_ATUALIZACAO_CHECKLIST"] = item.atualizado_em

    if not grouped:
        return []

    item_ids = [item_id for row in grouped.values() for item_id in row["ITEM_IDS"]]
    empenhos = []
    if item_ids:
        empenhos = db.query(models.Empenho).filter(
            models.Empenho.bom_item_id.in_(item_ids)
        ).all()

    total_por_item = {}
    for emp in empenhos:
        total_por_item[emp.bom_item_id] = total_por_item.get(emp.bom_item_id, 0.0) + float(emp.qtd_empenhada or 0)

    rows = []
    for row in grouped.values():
        qtd_lancada = sum(total_por_item.get(item_id, 0.0) for item_id in row["ITEM_IDS"])
        qtd_apontada = sum(
            resolve_bom_item_pointed_quantity(itens_por_id.get(item_id), total_por_item.get(item_id, 0.0))
            for item_id in row["ITEM_IDS"]
        )
        qtd_prevista_total = row["QTD_PREVISTA_TOTAL"]
        has_prevista = any(parse_quantity_value(v) is not None for v in row["QTD_PREVISTA_ORIGINAL"])
        qtd_prevista_value = qtd_prevista_total if has_prevista else None
        saldo = None if qtd_prevista_value is None else qtd_prevista_value - qtd_apontada
        rows.append(
            {
                "CHASSI": row["CHASSI"],
                "ORDEM_SERVICO_ARQUIVO": row["ORDEM_SERVICO_ARQUIVO"],
                "COD_ITEM": row["COD_ITEM"],
                "ITEM": row["ITEM"],
                "DESCRICAO": row["DESCRICAO"],
                "QTD_PREVISTA": format_quantity_value(qtd_prevista_value),
                "QTD_APONTADA": format_quantity_value(qtd_apontada),
                "QTD_CONSUMIDA": format_quantity_value(qtd_apontada),
                "QTD_LANCADA_HISTORICO": format_quantity_value(qtd_lancada),
                "SALDO": format_quantity_value(saldo),
                "STATUS_CONSUMO": compute_consumption_status(qtd_prevista_value, qtd_apontada),
                "STATUS_CHECKLIST": row["STATUS_CHECKLIST"],
                "RESPONSAVEL_CHECKLIST": ", ".join(sorted(set(filter(None, row["RESPONSAVEL_CHECKLIST"])))),
                "ATUALIZADO_EM": to_excel_dt(row["ULTIMA_ATUALIZACAO_CHECKLIST"]),
            }
        )
    return rows


def normalize_lookup_key(value: str) -> str:
    text = safe_str(value)
    if not text:
        return ""
    text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    text = re.sub(r"[^A-Z0-9]+", "_", text.upper()).strip("_")
    return text


def normalize_status_value(value: str) -> str:
    raw = safe_str(value)
    if not raw:
        return "N/A"
    lookup = normalize_lookup_key(raw)
    if lookup in {"SIM", "S", "OK"}:
        return "SIM"
    if lookup in {"NAO", "N", "X"}:
        return "NAO"
    if lookup == "N_A":
        return "N/A"
    return raw.upper()

def normalize_etapa(value: str) -> str:
    if not value:
        return ""
    v = str(value).strip().upper()
    v = v.replace("  ", " ")

    if v in ["AC", "A/C"]:
        return "A/C"
    if v in ["LIBERA", "LIBERA."]:
        return "LIBERA."
    if v in ["ACESSO", "ACESSO.", "ACESSÓ", "ACESSÓ."]:
        return "ACESSÓ."
    if v in ["SERRA", "SERRA."]:
        return "SERRA."
    if v in ["DESMON", "DESMONT"]:
        return "DESMONT"
    if v in ["ELETRICA", "ELÉTRICA", "ELÉTRIC", "ELÉTRIC."]:
        return "ELÉTRICA"
    return v


def is_done_status(value: str) -> bool:
    return str(value or "").strip().upper() in ["SIM", "S", "OK", "N/A"]


def is_pending_status(value: str) -> bool:
    return str(value or "").strip().upper() in ["N", "NÃO", "NAO", "X"]

# Define regras de filtragem por etapa
# Ajustado para validar contra "NÃO" e "SIM" conforme consta no banco de dados
ETAPA_REGRAS = {
    "VIDROS": lambda s: is_pending_status(s.get("VIDROS")),
    "A/C": lambda s: is_pending_status(s.get("A/C")),
    "PREP": lambda s: is_pending_status(s.get("PREP")),
    "SERRA.": lambda s: is_pending_status(s.get("SERRA.")),
    "EXPE.": lambda s: is_pending_status(s.get("EXPE.")),
    "DESMONT": lambda s: is_done_status(s.get("VIDROS")) and is_done_status(s.get("A/C")) and is_pending_status(s.get("DESMONT")),
    "ELÉTRICA": lambda s: is_done_status(s.get("DESMONT")) and is_pending_status(s.get("ELÉTRICA")),
    "REVEST": lambda s: is_done_status(s.get("DESMONT")) and is_pending_status(s.get("REVEST")),
    "BCO": lambda s: is_done_status(s.get("REVEST")) and is_pending_status(s.get("BCO")),
    "ACESSÓ.": lambda s: is_pending_status(s.get("ACESSÓ.")),
    "PLOTA.": lambda s: is_pending_status(s.get("PLOTA.")),
    "LIBERA.": lambda s: is_done_status(s.get("BCO")) and is_pending_status(s.get("LIBERA."))
}

AUTOMATIC_SEQUENCE_STAGE_TO_POSTOS = [
    ("VIDROS", ["CORTE_VIDRO"]),
    ("DESMONT", ["DESMONTAGEM"]),
    ("ELETRICA", ["ELETRICA"]),
    ("REVEST", ["REVESTIMENTO_1", "REVESTIMENTO_2", "REVESTIMENTO_3", "REVESTIMENTO_4"]),
    ("BCO", ["BANCOS"]),
    ("PREP", ["PREPARACAO"]),
    ("SERRA", ["TERCERIZACAO"]),
    ("EXPE", ["EXPEDICAO"]),
    ("LIBERA", ["LIBERACAO"]),
]


def build_status_map(aponts):
    return {
        normalize_etapa(a.etapa): str(a.status).strip().upper()
        for a in aponts
    }


def apply_vehicle_progress(veiculo, status_map: dict):
    concluidos = sum(
        1 for e in ETAPAS_PRODUCAO
        if status_map.get(e.upper()) in ["SIM", "S", "OK", "N/A"]
    )
    veiculo.progresso = int((concluidos / len(ETAPAS_PRODUCAO)) * 100) if ETAPAS_PRODUCAO else 0

    veiculo.etapa_atual = "FINALIZADO"
    for e in ETAPAS_STATUS_ATUAL:
        if status_map.get(e.upper()) not in ["SIM", "S", "OK", "N/A"]:
            veiculo.etapa_atual = e
            break


def vehicle_matches_stage_filter(veiculo, filtro: str, status_map: dict) -> bool:
    filtro_limpo = str(filtro or "").strip().upper()
    if not filtro_limpo:
        return True

    status_map_normalizado = {
        normalize_etapa(k): str(v).strip().upper()
        for k, v in (status_map or {}).items()
    }

    if filtro_limpo in ["GE", "CLIM"]:
        return (veiculo.ar_condicionado or "").strip().upper() == filtro_limpo and ETAPA_REGRAS["A/C"](status_map_normalizado)

    filtro_norm = normalize_etapa(filtro_limpo)
    if filtro_norm == "BCO":
        banco_flag = (veiculo.banco_presente or "").strip().upper()
        if banco_flag in ["N", "NAO", "NÃO", "NAO TEM", "SEM", "0"]:
            return False

    regra = ETAPA_REGRAS.get(filtro_norm)
    return bool(regra and regra(status_map_normalizado))


def get_apontamentos_by_chassi(db: Session, veiculos_db):
    chassis = [str(v.chassi).strip() for v in veiculos_db]
    apontamentos = []
    if chassis:
        apontamentos = db.query(models.Apontamento).filter(
            func.trim(cast(models.Apontamento.chassi, String)).in_(chassis)
        ).all()

    apont_por_chassi = {}
    for apontamento in apontamentos:
        ch_key = str(apontamento.chassi).strip()
        apont_por_chassi.setdefault(ch_key, []).append(apontamento)
    return apont_por_chassi


def generate_automatic_sequences(db: Session):
    veiculos_db = active_vehicle_query(db).order_by(models.Veiculo.ordem.asc()).all()
    apont_por_chassi = get_apontamentos_by_chassi(db, veiculos_db)

    db.query(models.PostoSequencia).delete()
    db.flush()

    sequence_counters = {
        posto_key: 0
        for _, postos in AUTOMATIC_SEQUENCE_STAGE_TO_POSTOS
        for posto_key in postos
    }

    for stage_key, postos in AUTOMATIC_SEQUENCE_STAGE_TO_POSTOS:
        elegiveis = []
        for veiculo in veiculos_db:
            chassi_key = str(veiculo.chassi).strip()
            status_map = build_status_map(apont_por_chassi.get(chassi_key, []))
            if vehicle_matches_stage_filter(veiculo, stage_key, status_map):
                elegiveis.append(veiculo)

        for idx, veiculo in enumerate(elegiveis):
            posto_key = postos[idx % len(postos)]
            sequence_counters[posto_key] += 1
            db.add(
                models.PostoSequencia(
                    posto=posto_key,
                    chassi=str(veiculo.chassi).strip(),
                    sequencia=sequence_counters[posto_key],
                )
            )


DEMO_MODELOS = [
    "TORINO",
    "APACHE VIP",
    "ATTIVI",
    "MILLENNIUM",
    "VIALE",
    "IDEALE",
]

DEMO_CLIENTES = [
    "VIACAO ALPHA",
    "TRANS BETA",
    "URBANO SUL",
    "EXPRESSO NORTE",
    "MOBILIDADE OESTE",
]

DEMO_DESTINOS = [
    "SAO PAULO",
    "CURITIBA",
    "CAMPINAS",
    "SOROCABA",
    "RIBEIRAO PRETO",
]

DEMO_BANCO_COMENTARIOS = [
    "MDBUS",
    "INCORPOL",
    "AJUSTE FINAL",
    "LIBERADO",
    "AGUARDANDO KIT",
]

DEMO_RESPONSAVEIS = [
    "Carlos",
    "Marina",
    "Felipe",
    "Julia",
    "Renato",
    "Patricia",
]

DEMO_BOM_DESCRICOES = [
    ("1001", "PARAFUSO", "Parafuso de fixacao"),
    ("1002", "SUPORTE", "Suporte lateral"),
    ("1003", "CHICOTE", "Chicote eletrico"),
    ("1004", "ACABAMENTO", "Acabamento interno"),
    ("1005", "KIT PORTA", "Kit de porta"),
    ("1006", "LUMINARIA", "Luminaria interna"),
]


def reset_ordens_servico(db: Session) -> int:
    ordens_servico = db.query(models.OrdemServico).all()
    for ordem_servico in ordens_servico:
        remove_ordem_servico_arquivo(ordem_servico)
    removidas = len(ordens_servico)
    db.query(models.OrdemServico).delete()
    db.flush()
    return removidas


def clear_active_base(db: Session, preserve_ordens_servico: bool = False):
    db.query(models.Empenho).delete()
    db.query(models.BomItem).delete()
    if not preserve_ordens_servico:
        reset_ordens_servico(db)
    db.query(models.PostoSequencia).delete()
    db.query(models.Apontamento).delete()
    db.query(models.Veiculo).delete()
    db.flush()


def prepare_base_import(db: Session):
    db.query(models.PostoSequencia).delete()
    db.query(models.Veiculo).update(
        {"ativo": False},
        synchronize_session=False,
    )
    db.flush()


def get_latest_bom_responsavel(db: Session, tipo: str, chassi: str) -> str:
    item = db.query(models.BomItem).filter(
        models.BomItem.tipo == str(tipo).strip().upper(),
        func.trim(cast(models.BomItem.chassi, String)) == str(chassi).strip(),
    ).order_by(models.BomItem.atualizado_em.desc(), models.BomItem.id.desc()).first()
    return safe_str(item.responsavel if item else "") or "Sistema"


def build_manual_vehicle_stage_statuses(banco_presente: str) -> dict:
    status_map = {etapa: "NAO" for etapa in ETAPAS_PRODUCAO}
    banco_flag = safe_str(banco_presente).upper()
    if banco_flag in {"N", "NAO", "NÃO", "SEM", "0"}:
        status_map["BCO"] = "N/A"
    return status_map


def map_stage_status_from_raw_value(value) -> str:
    raw = safe_str(value).upper()
    if raw in {"S", "SIM", "OK"}:
        return "SIM"
    if raw in {"N", "NAO", "NÃO", "X"}:
        return "NAO"
    return "N/A"


def build_demo_status_map(
    main_stage: str,
    has_banco: bool,
    ac_pending: bool,
    prep_pending: bool,
    serra_pending: bool,
    expe_pending: bool,
):
    stage_lookup = {normalize_etapa(etapa): etapa for etapa in ETAPAS_PRODUCAO}
    
    def resolve_stage_key(alias: str) -> str:
        alias_norm = normalize_etapa(alias)
        if alias_norm in stage_lookup:
            return alias_norm
        for candidate in stage_lookup:
            if candidate.startswith(alias_norm) or alias_norm.startswith(candidate):
                return candidate
        raise KeyError(alias)

    stage_vidros = resolve_stage_key("VIDROS")
    stage_ac = resolve_stage_key("A/C")
    stage_prep = resolve_stage_key("PREP")
    stage_serra = resolve_stage_key("SERRA")
    stage_expe = resolve_stage_key("EXPE")
    stage_desmont = resolve_stage_key("DESMONT")
    stage_eletrica = resolve_stage_key("ELETRICA")
    stage_revest = resolve_stage_key("REVEST")
    stage_bco = resolve_stage_key("BCO")
    stage_acesso = resolve_stage_key("ACESSO")
    stage_plota = resolve_stage_key("PLOTA")
    stage_libera = resolve_stage_key("LIBERA")

    statuses = {etapa: "NAO" for etapa in ETAPAS_PRODUCAO}
    if not has_banco:
        statuses[stage_lookup[stage_bco]] = "N/A"

    if main_stage == "FINALIZADO":
        for etapa in ETAPAS_PRODUCAO:
            statuses[etapa] = "SIM"
        if not has_banco:
            statuses[stage_lookup[stage_bco]] = "N/A"
        return statuses

    done_by_main_stage = {
        stage_vidros: [],
        stage_desmont: [stage_vidros],
        stage_eletrica: [stage_vidros, stage_desmont],
        stage_revest: [stage_vidros, stage_desmont, stage_eletrica],
        stage_bco: [stage_vidros, stage_desmont, stage_eletrica, stage_revest],
        stage_libera: [stage_vidros, stage_desmont, stage_eletrica, stage_revest, stage_acesso, stage_plota],
    }

    for stage_norm in done_by_main_stage.get(main_stage, []):
        statuses[stage_lookup[stage_norm]] = "SIM"
    if main_stage == stage_libera and has_banco:
        statuses[stage_lookup[stage_bco]] = "SIM"

    statuses[stage_lookup[stage_ac]] = "NAO" if ac_pending else "SIM"
    statuses[stage_lookup[stage_prep]] = "NAO" if prep_pending else "SIM"
    statuses[stage_lookup[stage_serra]] = "NAO" if serra_pending else "SIM"
    statuses[stage_lookup[stage_expe]] = "NAO" if expe_pending else "SIM"

    if main_stage in stage_lookup:
        statuses[stage_lookup[main_stage]] = "NAO"

    if main_stage != stage_libera:
        statuses[stage_lookup[stage_libera]] = "NAO"
        statuses[stage_lookup[stage_acesso]] = "NAO"
        statuses[stage_lookup[stage_plota]] = "NAO"

    return statuses


def create_demo_ordem_servico(chassi: str, modelo: str, cliente: str, destino: str):
    if DocxDocument is None:
        return None

    caminho = OS_UPLOAD_DIR / f"demo_{chassi}.docx"
    doc = DocxDocument()
    doc.add_heading(f"Ordem de Servico - {chassi}", level=1)
    doc.add_paragraph(f"Modelo: {modelo}")
    doc.add_paragraph(f"Cliente: {cliente}")
    doc.add_paragraph(f"Destino: {destino}")
    doc.add_paragraph("Checklist sugerido para teste:")

    tabela = doc.add_table(rows=1, cols=3)
    cabecalho = tabela.rows[0].cells
    cabecalho[0].text = "Item"
    cabecalho[1].text = "Conferencia"
    cabecalho[2].text = "Observacao"

    for item in ["Estrutura", "Acabamento", "Eletrica"]:
        linha = tabela.add_row().cells
        linha[0].text = item
        linha[1].text = "Pendente"
        linha[2].text = "Gerado automaticamente"

    doc.save(caminho)
    return models.OrdemServico(
        chassi=chassi,
        nome_arquivo=f"O.S._{chassi}.docx",
        caminho_arquivo=str(caminho),
    )


def create_demo_bom_items(db: Session, chassi: str, tipo: str, done_count: int, rng: random.Random):
    selected_items = rng.sample(DEMO_BOM_DESCRICOES, k=min(4, len(DEMO_BOM_DESCRICOES)))
    created_items = []

    for idx, (codigo, item_nome, descricao) in enumerate(selected_items, start=1):
        quantidade = rng.randint(1, 6)
        item = models.BomItem(
            tipo=tipo,
            chassi=chassi,
            cod_item=f"{codigo}-{idx}",
            item=item_nome,
            descricao=descricao,
            qtd=str(quantidade),
            status="SIM" if idx <= done_count else "NAO",
        )
        db.add(item)
        created_items.append(item)

    db.flush()
    return created_items


def generate_demo_dataset(db: Session, quantidade: int = 24):
    rng = random.Random()
    agora = datetime.datetime.now(LOCAL_TZ)
    main_stages = [
        normalize_etapa("VIDROS"),
        normalize_etapa("DESMONT"),
        normalize_etapa("ELETRICA"),
        normalize_etapa("REVEST"),
        normalize_etapa("BCO"),
        normalize_etapa("LIBERA"),
        "FINALIZADO",
    ]
    stage_index_lookup = {
        normalize_etapa(etapa): idx
        for idx, etapa in enumerate(ETAPAS_PRODUCAO)
    }

    scenario_pool = []
    while len(scenario_pool) < quantidade:
        scenario_pool.extend(main_stages)
    rng.shuffle(scenario_pool)

    clear_active_base(db, preserve_ordens_servico=False)

    demo_rows = []
    for idx in range(quantidade):
        chassi = f"TESTE{idx + 1:05d}"
        modelo = rng.choice(DEMO_MODELOS)
        cliente = rng.choice(DEMO_CLIENTES)
        destino = rng.choice(DEMO_DESTINOS)
        ar_condicionado = "GE" if idx % 2 == 0 else "CLIM"
        has_banco = idx % 5 != 0
        banco_presente = "SIM" if has_banco else "NAO"
        banco_comentario = rng.choice(DEMO_BANCO_COMENTARIOS) if has_banco else "SEM BANCO"
        localizacao = rng.choice(LOCALIZACOES)
        main_stage = scenario_pool[idx]
        ac_pending = idx % 6 == 0
        prep_pending = idx % 3 == 0
        serra_pending = idx % 4 == 0
        expe_pending = idx % 5 in {0, 1}
        status_map = build_demo_status_map(
            main_stage=main_stage,
            has_banco=has_banco,
            ac_pending=ac_pending,
            prep_pending=prep_pending,
            serra_pending=serra_pending,
            expe_pending=expe_pending,
        )

        veiculo = models.Veiculo(
            chassi=chassi,
            modelo=modelo,
            ordem=idx + 1,
            ar_condicionado=ar_condicionado,
            cj_bco=f"CJ-{rng.randint(100, 999)}",
            cliente=cliente,
            destino=destino,
            localizacao=localizacao,
            banco_presente=banco_presente,
            banco_comentario=banco_comentario,
        )
        db.add(veiculo)

        for etapa in ETAPAS_PRODUCAO:
            status = status_map.get(etapa, "NAO")
            inicio = None
            termino = None
            responsavel = ""
            observacao = ""
            etapa_idx = stage_index_lookup.get(normalize_etapa(etapa), 0)
            inicio_base = agora - datetime.timedelta(days=max(0, quantidade - idx) // 5, hours=(len(ETAPAS_PRODUCAO) - etapa_idx))

            if status == "SIM":
                inicio = inicio_base
                termino = inicio + datetime.timedelta(minutes=35 + ((idx + etapa_idx) % 90))
                responsavel = rng.choice(DEMO_RESPONSAVEIS)
                observacao = "Concluido pela base de teste."
            elif normalize_etapa(etapa) == main_stage or (status == "NAO" and rng.random() < 0.12):
                inicio = inicio_base
                responsavel = rng.choice(DEMO_RESPONSAVEIS)
                observacao = "Em aberto na base de teste."

            db.add(
                models.Apontamento(
                    chassi=chassi,
                    etapa=etapa,
                    status=status,
                    responsavel=responsavel,
                    inicio=inicio,
                    termino=termino,
                    localizacao=localizacao,
                    observacao=observacao,
                )
            )

            if inicio:
                registrar_historico_evento(
                    db,
                    veiculo,
                    etapa,
                    status,
                    responsavel,
                    inicio=inicio,
                    termino=termino,
                    observacao=observacao,
                )

        prep_done = main_stage in {
            normalize_etapa("ELETRICA"),
            normalize_etapa("REVEST"),
            normalize_etapa("BCO"),
            normalize_etapa("LIBERA"),
            "FINALIZADO",
        } and not prep_pending
        expe_done = main_stage in {normalize_etapa("LIBERA"), "FINALIZADO"} and not expe_pending

        demo_rows.append(
            {
                "chassi": chassi,
                "modelo": modelo,
                "cliente": cliente,
                "destino": destino,
                "prep_done": prep_done,
                "expe_done": expe_done,
                "add_os": idx < min(8, quantidade),
                "responsavel": rng.choice(DEMO_RESPONSAVEIS),
            }
        )

    db.flush()

    empenhos_pendentes = []
    for row in demo_rows:
        prep_done_count = 4 if row["prep_done"] else rng.randint(0, 2)
        exp_done_count = 4 if row["expe_done"] else rng.randint(0, 2)

        create_demo_bom_items(db, row["chassi"], "PREPARACAO", prep_done_count, rng)
        exp_items = create_demo_bom_items(db, row["chassi"], "EXPEDICAO", exp_done_count, rng)

        sync_stage_from_bom(db, "PREPARACAO", row["chassi"], row["responsavel"])
        sync_stage_from_bom(db, "EXPEDICAO", row["chassi"], row["responsavel"])

        if row["add_os"]:
            ordem_servico = create_demo_ordem_servico(
                chassi=row["chassi"],
                modelo=row["modelo"],
                cliente=row["cliente"],
                destino=row["destino"],
            )
            if ordem_servico:
                db.add(ordem_servico)

        for item in exp_items:
            if str(item.status).strip().upper() != "SIM":
                continue
            qtd_item = float(str(item.qtd).replace(",", "."))
            qtd_empenhada = qtd_item if rng.random() < 0.6 else max(1.0, qtd_item - 1.0)
            empenhos_pendentes.append(
                {
                    "item": item,
                    "qtd_empenhada": qtd_empenhada,
                    "responsavel": row["responsavel"],
                }
            )

    generate_automatic_sequences(db)
    db.flush()

    for empenho_data in empenhos_pendentes:
        item = empenho_data["item"]
        db.add(
            models.Empenho(
                bom_item_id=item.id,
                chassi=item.chassi,
                cod_item=item.cod_item,
                item=item.item,
                descricao=item.descricao,
                qtd_empenhada=empenho_data["qtd_empenhada"],
                sequencia_producao=get_sequence_number(db, "EXPEDICAO", item.chassi),
                responsavel=empenho_data["responsavel"],
            )
        )

    return quantidade

@app.get("/")
async def home(request: Request, db: Session = Depends(database.get_db), modelo: str = None, etapa: str = None):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not is_management_user(request):
        return RedirectResponse(url=get_operator_home_url(request), status_code=303)
    query = active_vehicle_query(db)

    # Filtragem por texto (Modelo, Chassi, Ar Condicionado, CJ. BCO, Localização)
    # Adicionado func.coalesce para evitar que valores NULL quebrem a busca LIKE
    if modelo and modelo.strip():
        termo = f"%{modelo.strip().upper()}%"
        query = query.filter(
            or_(
                func.upper(func.coalesce(cast(models.Veiculo.modelo, String), "")).like(termo),
                func.upper(func.coalesce(cast(models.Veiculo.chassi, String), "")).like(termo),
                func.upper(func.coalesce(cast(models.Veiculo.ar_condicionado, String), "")).like(termo),
                func.upper(func.coalesce(cast(models.Veiculo.cj_bco, String), "")).like(termo),
                func.upper(func.coalesce(cast(models.Veiculo.localizacao, String), "")).like(termo)
            )
        )

    veiculos_db = query.order_by(models.Veiculo.ordem.asc()).all()
    veiculos_exibicao = []
    apont_por_chassi = get_apontamentos_by_chassi(db, veiculos_db)

    for v in veiculos_db:
        chassi_key = str(v.chassi).strip()
        aponts = apont_por_chassi.get(chassi_key, [])

        status_map = build_status_map(aponts)
        apply_vehicle_progress(v, status_map)

        if etapa and etapa.strip():
            if vehicle_matches_stage_filter(v, etapa, status_map):
                veiculos_exibicao.append(v)
        else:
            veiculos_exibicao.append(v)

    return templates.TemplateResponse(
        request,
        "index.html",
        {
            "request": request,
            "veiculos": veiculos_exibicao,
            "etapas": ETAPAS_FILTRO,
            "termo_busca": modelo or "",
            "etapa_selecionada": etapa or "",
            "current_user": require_login(request),
            "is_admin": bool(require_admin(request)),
        }
    )

@app.get("/veiculo/{chassi}")
async def detalhes(request: Request, chassi: str, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not is_management_user(request):
        return RedirectResponse(url=get_operator_home_url(request), status_code=303)
    c_limpo = chassi.strip()
    user_name = get_user_name(request)

    veiculo = get_vehicle_by_chassi(db, c_limpo)
    if not veiculo:
        return RedirectResponse(url="/", status_code=303)

    feitos = db.query(models.Apontamento).filter(
        func.trim(cast(models.Apontamento.chassi, String)) == c_limpo
    ).all()

    for f in feitos:
        f.inicio_str = to_input_dt(f.inicio)
        f.termino_str = to_input_dt(f.termino)

    apont_map = {
        normalize_etapa(f.etapa): f
        for f in feitos
    }
    ordem_servico = get_ordem_servico(db, c_limpo)
    etapas_bloqueadas_gestao = get_management_locked_stages_for_profile(get_user_profile(request))

    return templates.TemplateResponse(
        request,
        "detalhes.html",
        {
            "request": request,
            "veiculo": veiculo,
            "etapas": ETAPAS_PRODUCAO,
            "apont_map": apont_map,
            "localizacoes": LOCALIZACOES,
            "user_name": user_name,
            "current_user": require_login(request),
            "is_admin": bool(require_admin(request)),
            "ordem_servico": ordem_servico,
            "etapas_bloqueadas_gestao": etapas_bloqueadas_gestao,
        }
    )

@app.post("/upload")
async def upload_base(request: Request, file: UploadFile = File(...), db: Session = Depends(database.get_db)):
    # Protege upload com login simples
    # (mantém compatível sem usuários cadastrados)
    if not require_login(request):
        return {"status": "erro", "detail": "Login necessário"}
    if not require_admin(request):
        return {"status": "erro", "detail": "Acesso restrito ao ADM"}
    try:
        content = await file.read()

        df = (
            pd.read_excel(io.BytesIO(content))
            if file.filename.endswith(".xlsx")
            else pd.read_csv(io.BytesIO(content))
        )

        df.columns = [str(c).upper().strip() for c in df.columns]

        def get_col(row, *names):
            for n in names:
                if n in df.columns:
                    val = row.get(n, "")
                    if pd.isna(val):
                        return ""
                    return str(val).strip()
            return ""

        etapas_col = {normalize_etapa(c): c for c in df.columns}

        # Prepara nova carga sem apagar O.S., B.O.M. e empenhos.
        prepare_base_import(db)

        imported_rows = []
        imported_chassis = []
        for idx, row in df.iterrows():
            ch_raw = str(row.get("CHASSI", "")).strip().split(".")[0]
            if not ch_raw or ch_raw.lower() == "nan":
                continue

            imported_rows.append(
                {
                    "chassi": ch_raw,
                    "ordem": int(idx) + 1,
                    "modelo": str(row.get("MMMV", "")).strip().upper(),
                    "ar_condicionado": get_col(row, "AR CONDICIONADO", "AR_CONDICIONADO", "AR-CONDICIONADO", "ARCONDICIONADO"),
                    "cj_bco": get_col(row, "CJ. BCO", "CJ BCO", "CJ_BCO", "CJ-BCO"),
                    "cliente": get_col(row, "CLIENTE"),
                    "destino": get_col(row, "DESTINO"),
                    "localizacao": get_col(row, "LOCALIZACAO", "LOCALIZAÇÃO"),
                    "banco_presente": get_col(row, "BANCO", "BANCO_PRESENTE", "POSSUI BANCO", "TEM BANCO"),
                    "banco_comentario": get_col(row, "COMENTARIO BANCO", "COMENTARIO_BANCO", "BANCO OBS", "OBS BANCO"),
                    "row": row,
                }
            )
            imported_chassis.append(ch_raw)

        if not imported_rows:
            return {"status": "erro", "detail": "Nenhum chassi valido encontrado na base."}

        db.query(models.Apontamento).filter(
            func.trim(cast(models.Apontamento.chassi, String)).in_(imported_chassis)
        ).delete(synchronize_session=False)

        existing_vehicles = {
            str(item.chassi).strip(): item
            for item in db.query(models.Veiculo).filter(
                func.trim(cast(models.Veiculo.chassi, String)).in_(imported_chassis)
            ).all()
        }

        for item in imported_rows:
            veiculo = existing_vehicles.get(item["chassi"])
            if not veiculo:
                veiculo = models.Veiculo(chassi=item["chassi"])
                db.add(veiculo)
                existing_vehicles[item["chassi"]] = veiculo

            veiculo.modelo = item["modelo"]
            veiculo.ordem = item["ordem"]
            veiculo.ativo = True
            veiculo.ar_condicionado = item["ar_condicionado"]
            veiculo.cj_bco = item["cj_bco"]
            veiculo.cliente = item["cliente"]
            veiculo.destino = item["destino"]
            veiculo.localizacao = item["localizacao"]
            veiculo.banco_presente = item["banco_presente"]
            veiculo.banco_comentario = item["banco_comentario"]

            for etapa in ETAPAS_PRODUCAO:
                col_name = etapas_col.get(normalize_etapa(etapa))
                if col_name:
                    status = map_stage_status_from_raw_value(item["row"][col_name])
                else:
                    status = "N/A"
                db.add(
                    models.Apontamento(
                        chassi=item["chassi"],
                        etapa=etapa,
                        status=status
                    )
                )

        db.flush()

        for chassi in imported_chassis:
            for tipo in ["PREPARACAO", "EXPEDICAO"]:
                if get_bom_items(db, tipo, chassi):
                    sync_stage_from_bom(db, tipo, chassi, get_latest_bom_responsavel(db, tipo, chassi))

        db.commit()
        return {"status": "sucesso"}

    except Exception as e:
        db.rollback()
        return {"status": "erro", "detail": str(e)}


@app.post("/gerar_base_teste")
async def gerar_base_teste(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return {"status": "erro", "detail": "Login necessario"}
    if not require_admin(request):
        return {"status": "erro", "detail": "Acesso restrito ao ADM"}

    try:
        quantidade = generate_demo_dataset(db)
        db.commit()
        return {
            "status": "sucesso",
            "detail": f"{quantidade} registros de teste gerados com sucesso.",
        }
    except Exception as exc:
        db.rollback()
        return {"status": "erro", "detail": str(exc)}

@app.post("/upload_apontamentos")
async def upload_apontamentos(request: Request, file: UploadFile = File(...), db: Session = Depends(database.get_db)):
    if not require_login(request):
        return {"status": "erro", "detail": "Login necessário"}
    if not require_admin(request):
        return {"status": "erro", "detail": "Acesso restrito ao ADM"}
    try:
        content = await file.read()

        df = (
            pd.read_excel(io.BytesIO(content))
            if file.filename.endswith(".xlsx")
            else pd.read_csv(io.BytesIO(content))
        )

        df.columns = [str(c).upper().strip() for c in df.columns]

        # Normaliza e agrega para evitar N+1
        rows = {}
        banco_updates = {}

        for _, row in df.iterrows():
            ch_raw = safe_str(row.get("CHASSI", "")).split(".")[0]
            if not ch_raw or ch_raw.lower() == "nan":
                continue

            etapa = normalize_etapa(safe_str(row.get("ETAPA", "")))
            inicio = parse_local_dt(row.get("INICIO"))
            termino = parse_local_dt(row.get("TERMINO"))
            responsavel = safe_str(row.get("RESPONSAVEL", ""))

            banco_presente = safe_str(row.get("BANCO", ""))
            banco_comentario = safe_str(row.get("COMENTARIO BANCO", row.get("COMENTARIO_BANCO", "")))
            if banco_presente or banco_comentario:
                banco_updates[ch_raw] = {
                    "banco_presente": banco_presente,
                    "banco_comentario": banco_comentario
                }

            if not etapa:
                continue

            rows[(ch_raw, etapa)] = {
                "inicio": inicio,
                "termino": termino,
                "responsavel": responsavel
            }

        # Atualiza banco/comentário em lote
        for ch_raw, data in banco_updates.items():
            update_data = {}
            if data.get("banco_presente"):
                update_data["banco_presente"] = data["banco_presente"]
            if data.get("banco_comentario"):
                update_data["banco_comentario"] = data["banco_comentario"]
            if update_data:
                db.query(models.Veiculo).filter(
                    func.trim(cast(models.Veiculo.chassi, String)) == ch_raw,
                    active_vehicle_filter(),
                ).update(update_data)

        if rows:
            chassis = list({k[0] for k in rows.keys()})
            existentes = db.query(models.Apontamento).filter(
                func.trim(cast(models.Apontamento.chassi, String)).in_(chassis)
            ).all()
            existentes_map = {
                (str(a.chassi).strip(), normalize_etapa(a.etapa)): a
                for a in existentes
            }

            for (ch_raw, etapa), data in rows.items():
                ap = existentes_map.get((ch_raw, etapa))
                if not ap:
                    ap = models.Apontamento(
                        chassi=ch_raw,
                        etapa=etapa,
                        status="N/A"
                    )
                    db.add(ap)
                ap.inicio = data["inicio"]
                ap.termino = data["termino"]
                ap.responsavel = data["responsavel"]

        db.commit()
        return {"status": "sucesso"}

    except Exception as e:
        db.rollback()
        return {"status": "erro", "detail": str(e)}

@app.post("/apontar")
async def salvar(request: Request, data: dict = Body(...), db: Session = Depends(database.get_db)):
    user = require_login(request)
    if not user:
        return {"status": "erro", "detail": "Login necessário"}
    if not is_management_user(request):
        return {"status": "erro", "detail": "Ação disponível apenas para ADM e LIDER"}

    ch = str(data["chassi"]).strip()
    et = normalize_etapa(data["etapa"])
    etapas_bloqueadas_gestao = get_management_locked_stages_for_profile(user.get("perfil", ""))
    if et in etapas_bloqueadas_gestao:
        return {"status": "erro", "detail": etapas_bloqueadas_gestao[et]}

    st = normalize_status_value(data.get("status", ""))
    responsavel = str(data.get("responsavel", "")).strip()
    inicio = parse_local_dt(data.get("inicio"))
    termino = parse_local_dt(data.get("termino"))

    registrar_historico = bool(data.get("registrar_historico", True))

    # Atualiza ou cria o apontamento
    db.query(models.Apontamento).filter(
        func.trim(cast(models.Apontamento.chassi, String)) == ch,
        func.trim(cast(models.Apontamento.etapa, String)) == et
    ).delete()

    db.add(models.Apontamento(
        chassi=ch,
        etapa=et,
        status=st,
        responsavel=responsavel,
        inicio=inicio,
        termino=termino,
        localizacao=None
    ))

    v = get_vehicle_by_chassi(db, ch)

    # Registra no histórico apenas quando for status (SIM/NÃO/N/A) e explícito
    if registrar_historico:
        db.add(models.Historico(
            chassi=ch,
            modelo=v.modelo if v else "N/A",
            etapa=et,
            status=st,
            responsavel=responsavel,
            inicio=inicio,
            termino=termino,
            localizacao=None
        ))

    db.commit()
    return {"status": "ok"}

@app.post("/veiculo_localizacao")
async def atualizar_localizacao(data: dict = Body(...), db: Session = Depends(database.get_db)):
    ch = str(data.get("chassi", "")).strip()
    localizacao = str(data.get("localizacao", "")).strip()
    if not ch:
        return {"status": "erro", "detail": "Chassi inválido"}

    db.query(models.Veiculo).filter(
        func.trim(cast(models.Veiculo.chassi, String)) == ch,
        active_vehicle_filter(),
    ).update({"localizacao": localizacao})
    db.commit()
    return {"status": "ok"}

@app.post("/veiculo_banco")
async def atualizar_banco(data: dict = Body(...), db: Session = Depends(database.get_db)):
    ch = str(data.get("chassi", "")).strip()
    banco_presente = str(data.get("banco_presente", "")).strip()
    banco_comentario = str(data.get("banco_comentario", "")).strip()
    if not ch:
        return {"status": "erro", "detail": "Chassi inválido"}

    db.query(models.Veiculo).filter(
        func.trim(cast(models.Veiculo.chassi, String)) == ch,
        active_vehicle_filter(),
    ).update({
        "banco_presente": banco_presente,
        "banco_comentario": banco_comentario
    })
    db.commit()
    return {"status": "ok"}

@app.get("/exportar_historico")
async def exportar(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)
    logs = db.query(models.Historico).all()
    if not logs:
        return {"message": "Sem dados"}

    veiculos = db.query(models.Veiculo).all()
    loc_map = {str(v.chassi).strip(): v.localizacao for v in veiculos}
    apont_map = {}
    aponts = db.query(models.Apontamento).all()
    for a in aponts:
        apont_map[(str(a.chassi).strip(), normalize_etapa(a.etapa))] = a

    df = pd.DataFrame([
        {
            "CHASSI": l.chassi,
            "MODELO": l.modelo,
            "ETAPA": normalize_etapa(l.etapa),
            "STATUS": l.status,
            "RESPONSAVEL": (apont_map.get((str(l.chassi).strip(), normalize_etapa(l.etapa))) or l).responsavel,
            "INICIO": to_excel_dt((apont_map.get((str(l.chassi).strip(), normalize_etapa(l.etapa))) or l).inicio),
            "TERMINO": to_excel_dt((apont_map.get((str(l.chassi).strip(), normalize_etapa(l.etapa))) or l).termino),
            "LOCALIZACAO": loc_map.get(str(l.chassi).strip()),
            "DATA": to_excel_dt(l.data_apontamento)
        }
        for l in logs
    ])

    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as w:
        df.to_excel(w, index=False)

    out.seek(0)

    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=relatorio.xlsx"}
    )

@app.get("/exportar_tempos")
async def exportar_tempos(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)
    aponts = db.query(models.Apontamento).all()
    veiculos = db.query(models.Veiculo).all()
    loc_map = {str(v.chassi).strip(): v.localizacao for v in veiculos}
    modelo_map = {str(v.chassi).strip(): v.modelo for v in veiculos}

    df = pd.DataFrame([
        {
            "CHASSI": a.chassi,
            "MODELO": modelo_map.get(str(a.chassi).strip()),
            "ETAPA": normalize_etapa(a.etapa),
            "RESPONSAVEL": a.responsavel,
            "INICIO": to_excel_dt(a.inicio),
            "TERMINO": to_excel_dt(a.termino),
            "LOCALIZACAO": loc_map.get(str(a.chassi).strip())
        }
        for a in aponts
    ])

    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as w:
        df.to_excel(w, index=False)

    out.seek(0)

    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=tempos_localizacao.xlsx"}
    )


@app.get("/exportacoes", response_class=HTMLResponse)
async def exportacoes_page(request: Request):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)
    return templates.TemplateResponse(
        request,
        "exportacoes.html",
        {
            "request": request,
            "current_user": require_login(request),
        },
    )


@app.get("/modelo_bom_padrao")
async def exportar_modelo_bom_padrao(request: Request):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)

    df = pd.DataFrame(
        columns=["nº chassi", "cod. item", "item", "descrição", "qtd"]
    )
    return dataframe_to_excel_response(df, "modelo_bom_padrao.xlsx")


@app.get("/exportar_veiculos")
async def exportar_veiculos(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)

    veiculos = active_vehicle_query(db).order_by(models.Veiculo.ordem.asc()).all()
    df = pd.DataFrame(
        [
            {
                "CHASSI": v.chassi,
                "MODELO": v.modelo,
                "ORDEM": v.ordem,
                "AR_CONDICIONADO": v.ar_condicionado,
                "CJ_BCO": v.cj_bco,
                "CLIENTE": v.cliente,
                "DESTINO": v.destino,
                "LOCALIZACAO": v.localizacao,
                "BANCO_PRESENTE": v.banco_presente,
                "BANCO_COMENTARIO": v.banco_comentario,
            }
            for v in veiculos
        ]
    )
    return dataframe_to_excel_response(df, "veiculos_base.xlsx")


@app.get("/exportar_sequenciamento")
async def exportar_sequenciamento(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)

    itens = db.query(models.PostoSequencia).order_by(models.PostoSequencia.posto.asc(), models.PostoSequencia.sequencia.asc()).all()
    df = pd.DataFrame(
        [
            {
                "POSTO": item.posto,
                "POSTO_LABEL": (get_posto_config(item.posto) or {}).get("label"),
                "CHASSI": item.chassi,
                "SEQUENCIA": item.sequencia,
            }
            for item in itens
        ]
    )
    return dataframe_to_excel_response(df, "sequenciamento_postos.xlsx")


@app.get("/exportar_bom_preparacao")
async def exportar_bom_preparacao(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)

    active_chassis = get_active_chassis(db)
    itens = db.query(models.BomItem).filter(
        models.BomItem.tipo == "PREPARACAO",
        models.BomItem.chassi.in_(active_chassis),
    ).order_by(models.BomItem.chassi.asc(), models.BomItem.id.asc()).all() if active_chassis else []
    resumo_por_item = build_bom_item_empenho_summary(itens, {})
    df = pd.DataFrame(
        [
            {
                "CHASSI": item.chassi,
                "COD_ITEM": item.cod_item,
                "ITEM": item.item,
                "DESCRICAO": item.descricao,
                "QTD": item.qtd,
                "QTD_APONTADA": resumo_por_item.get(item.id, {}).get("qtd_apontada_fmt"),
                "SALDO": resumo_por_item.get(item.id, {}).get("saldo_fmt"),
                "STATUS_CONSUMO": resumo_por_item.get(item.id, {}).get("status_consumo"),
                "STATUS": item.status,
                "RESPONSAVEL": item.responsavel,
                "ATUALIZADO_EM": to_excel_dt(item.atualizado_em),
            }
            for item in itens
        ]
    )
    return dataframe_to_excel_response(df, "bom_preparacao.xlsx")


@app.get("/exportar_bom_expedicao")
async def exportar_bom_expedicao(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not can_export_expedicao_lancamentos(request):
        return RedirectResponse(url="/", status_code=303)

    df = pd.DataFrame(build_expedicao_export_rows(db))
    return dataframe_to_excel_response(df, "bom_expedicao.xlsx")


@app.get("/exportar_empenhos")
async def exportar_empenhos(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not can_export_expedicao_lancamentos(request):
        return RedirectResponse(url="/", status_code=303)

    active_chassis = get_active_chassis(db)
    empenhos = db.query(models.Empenho).filter(
        models.Empenho.chassi.in_(active_chassis)
    ).order_by(models.Empenho.sequencia_producao.asc(), models.Empenho.criado_em.asc()).all() if active_chassis else []
    consolidated_rows = build_expedicao_export_rows(db)
    consolidated_map = {
        (
            safe_str(row["CHASSI"]),
            safe_str(row["COD_ITEM"]),
            safe_str(row["ITEM"]),
            safe_str(row["DESCRICAO"]),
        ): row
        for row in consolidated_rows
    }
    ordens_por_chassi = {
        ordem.chassi: ordem
        for ordem in db.query(models.OrdemServico).all()
    }
    df = pd.DataFrame(
        [
            {
                "CHASSI": item.chassi,
                "ORDEM_SERVICO_ARQUIVO": safe_str(ordens_por_chassi.get(item.chassi).nome_arquivo if ordens_por_chassi.get(item.chassi) else ""),
                "COD_ITEM": item.cod_item,
                "ITEM": item.item,
                "DESCRICAO": item.descricao,
                "QTD_LANCADA": format_quantity_value(float(item.qtd_empenhada or 0)),
                "QTD_CONSUMIDA_ITEM": consolidated_map.get(
                    (
                        safe_str(item.chassi),
                        safe_str(item.cod_item),
                        safe_str(item.item),
                        safe_str(item.descricao),
                    ),
                    {},
                ).get("QTD_CONSUMIDA"),
                "QTD_APONTADA_ITEM": consolidated_map.get(
                    (
                        safe_str(item.chassi),
                        safe_str(item.cod_item),
                        safe_str(item.item),
                        safe_str(item.descricao),
                    ),
                    {},
                ).get("QTD_APONTADA"),
                "QTD_PREVISTA_ITEM": consolidated_map.get(
                    (
                        safe_str(item.chassi),
                        safe_str(item.cod_item),
                        safe_str(item.item),
                        safe_str(item.descricao),
                    ),
                    {},
                ).get("QTD_PREVISTA"),
                "STATUS_CONSUMO_ITEM": consolidated_map.get(
                    (
                        safe_str(item.chassi),
                        safe_str(item.cod_item),
                        safe_str(item.item),
                        safe_str(item.descricao),
                    ),
                    {},
                ).get("STATUS_CONSUMO"),
                "SEQUENCIA_PRODUCAO": item.sequencia_producao,
                "RESPONSAVEL": item.responsavel,
                "CRIADO_EM": to_excel_dt(item.criado_em),
            }
            for item in empenhos
        ]
    )
    return dataframe_to_excel_response(df, "empenhos_expedicao.xlsx")


@app.get("/exportar_ordens_servico")
async def exportar_ordens_servico(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)

    ordens = db.query(models.OrdemServico).order_by(models.OrdemServico.chassi.asc()).all()
    df = pd.DataFrame(
        [
            {
                "CHASSI": ordem.chassi,
                "NOME_ARQUIVO": ordem.nome_arquivo,
                "CAMINHO_ARQUIVO": ordem.caminho_arquivo,
                "CRIADO_EM": to_excel_dt(ordem.criado_em),
            }
            for ordem in ordens
        ]
    )
    return dataframe_to_excel_response(df, "ordens_servico.xlsx")

@app.get("/limpar_historico")
async def limpar_logs(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)
    db.query(models.Historico).delete()
    db.commit()
    return RedirectResponse(url="/", status_code=303)


@app.post("/resetar_ordens_servico")
async def resetar_ordens_servico(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return {"status": "erro", "detail": "Login necessario"}
    if not require_admin(request):
        return {"status": "erro", "detail": "Acesso restrito ao ADM"}

    try:
        removidas = reset_ordens_servico(db)
        db.commit()
        return {
            "status": "sucesso",
            "detail": f"{removidas} ordens de servico removidas com sucesso.",
        }
    except Exception as exc:
        db.rollback()
        return {"status": "erro", "detail": str(exc)}


@app.post("/resetar_empenho_obsoleto")
async def resetar_empenho_obsoleto(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return {"status": "erro", "detail": "Login necessario"}
    if not require_admin(request):
        return {"status": "erro", "detail": "Acesso restrito ao ADM"}

    try:
        active_chassis = get_active_chassis(db)
        query = db.query(models.Empenho)
        if active_chassis:
            query = query.filter(~models.Empenho.chassi.in_(active_chassis))
        removidos = query.count()
        query.delete(synchronize_session=False)
        db.commit()
        return {
            "status": "sucesso",
            "detail": f"{removidos} empenhos obsoletos removidos com sucesso.",
        }
    except Exception as exc:
        db.rollback()
        return {"status": "erro", "detail": str(exc)}


@app.post("/cadastros/veiculos")
async def cadastrar_veiculo(
    request: Request,
    chassi: str = Form(...),
    modelo: str = Form(...),
    ordem: str = Form(""),
    ar_condicionado: str = Form(""),
    cj_bco: str = Form(""),
    cliente: str = Form(""),
    destino: str = Form(""),
    localizacao: str = Form(""),
    banco_presente: str = Form(""),
    banco_comentario: str = Form(""),
    db: Session = Depends(database.get_db),
):
    if not require_login(request):
        return {"status": "erro", "detail": "Login necessario"}
    if not require_admin(request):
        return {"status": "erro", "detail": "Acesso restrito ao ADM"}

    chassi_key = safe_str(chassi).split(".")[0]
    modelo_key = safe_str(modelo).upper()
    if not chassi_key or not modelo_key:
        return {"status": "erro", "detail": "Informe ao menos chassi e modelo."}

    try:
        ordem_key = int(str(ordem).strip()) if safe_str(ordem) else get_next_vehicle_order(db)
    except ValueError:
        return {"status": "erro", "detail": "A ordem precisa ser numerica."}

    try:
        veiculo = get_vehicle_by_chassi(db, chassi_key, active_only=False)
        criado = veiculo is None
        if not veiculo:
            veiculo = models.Veiculo(chassi=chassi_key)
            db.add(veiculo)

        veiculo.modelo = modelo_key
        veiculo.ordem = ordem_key
        veiculo.ativo = True
        veiculo.ar_condicionado = safe_str(ar_condicionado).upper()
        veiculo.cj_bco = safe_str(cj_bco)
        veiculo.cliente = safe_str(cliente)
        veiculo.destino = safe_str(destino)
        veiculo.localizacao = safe_str(localizacao)
        veiculo.banco_presente = safe_str(banco_presente).upper()
        veiculo.banco_comentario = safe_str(banco_comentario)

        status_map = build_manual_vehicle_stage_statuses(veiculo.banco_presente)
        for etapa in ETAPAS_PRODUCAO:
            apont = get_apontamento_for_stage(db, chassi_key, etapa)
            if not apont:
                db.add(
                    models.Apontamento(
                        chassi=chassi_key,
                        etapa=etapa,
                        status=status_map.get(etapa, "NAO"),
                        localizacao=veiculo.localizacao,
                    )
                )

        db.flush()
        for tipo in ["PREPARACAO", "EXPEDICAO"]:
            if get_bom_items(db, tipo, chassi_key):
                sync_stage_from_bom(db, tipo, chassi_key, get_latest_bom_responsavel(db, tipo, chassi_key))

        db.commit()
        mensagem = "Veiculo cadastrado com sucesso." if criado else "Veiculo atualizado com sucesso."
        return {"status": "sucesso", "detail": mensagem}
    except Exception as exc:
        db.rollback()
        return {"status": "erro", "detail": str(exc)}

@app.get("/postos")
async def postos_page(request: Request):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if is_management_user(request):
        return RedirectResponse(url="/", status_code=303)

    perfil = get_user_profile(request)
    if perfil in DIRECT_POST_BY_PROFILE:
        return RedirectResponse(url=f"/postos/{DIRECT_POST_BY_PROFILE[perfil]}", status_code=303)

    postos = [POSTOS_TRABALHO[key] | {"key": key} for key in get_allowed_posts_for_profile(perfil)]
    return templates.TemplateResponse(
        request,
        "postos.html",
        {
            "request": request,
            "current_user": require_login(request),
            "postos": postos,
        },
    )

@app.get("/postos/{posto}")
async def posto_cards(request: Request, posto: str, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if is_management_user(request):
        return RedirectResponse(url="/", status_code=303)

    posto_key = str(posto).strip().upper()
    posto_cfg = get_posto_config(posto_key)
    if not posto_cfg:
        return RedirectResponse(url="/postos", status_code=303)
    if not can_access_posto(request, posto_key):
        return RedirectResponse(url="/postos", status_code=303)

    return templates.TemplateResponse(
        request,
        "operacao_posto.html",
        {
            "request": request,
            "current_user": require_login(request),
            "posto": {"key": posto_key, **posto_cfg},
            "cards": get_posto_cards(db, posto_key),
        },
    )

@app.get("/postos/{posto}/{chassi}")
async def posto_card_detail(request: Request, posto: str, chassi: str, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if is_management_user(request):
        return RedirectResponse(url="/", status_code=303)

    posto_key = str(posto).strip().upper()
    posto_cfg = get_posto_config(posto_key)
    if not posto_cfg or not can_access_posto(request, posto_key):
        return RedirectResponse(url="/postos", status_code=303)

    veiculo = get_vehicle_by_chassi(db, chassi)
    if not veiculo:
        return RedirectResponse(url=f"/postos/{posto_key}", status_code=303)

    apont = get_apontamento_for_stage(db, veiculo.chassi, posto_cfg["etapa"])
    ordem_servico = get_ordem_servico(db, veiculo.chassi)
    if get_posto_mode(posto_key) == "checklist":
        bom_tipo = posto_cfg.get("bom_tipo")
        bom_itens = get_bom_items(db, bom_tipo, veiculo.chassi)
        empenhos_por_item = build_bom_item_empenho_summary(
            bom_itens,
            get_empenhos_for_items(db, bom_itens),
        )
        return templates.TemplateResponse(
            request,
            "operacao_checklist.html",
            {
                "request": request,
                "current_user": require_login(request),
                "posto": {"key": posto_key, **posto_cfg},
                "veiculo": veiculo,
                "apontamento": apont,
                "status_operacao": build_operacao_status(apont),
                "bom_itens": bom_itens,
                "ordem_servico": ordem_servico,
                "empenhos_por_item": empenhos_por_item,
                "permite_empenho": bool(posto_cfg.get("permite_empenho")),
            },
        )
    return templates.TemplateResponse(
        request,
        "operacao_detalhe.html",
        {
            "request": request,
            "current_user": require_login(request),
            "posto": {"key": posto_key, **posto_cfg},
            "veiculo": veiculo,
            "apontamento": apont,
            "status_operacao": build_operacao_status(apont),
            "ordem_servico": ordem_servico,
        },
    )

@app.post("/operacao/acao")
async def operacao_acao(request: Request, data: dict = Body(...), db: Session = Depends(database.get_db)):
    user = require_login(request)
    if not user:
        return {"status": "erro", "detail": "Login necessário"}
    if is_management_user(request):
        return {"status": "erro", "detail": "Ação disponível apenas na visão operacional"}

    posto_key = str(data.get("posto", "")).strip().upper()
    chassi = str(data.get("chassi", "")).strip()
    acao = str(data.get("acao", "")).strip().lower()
    observacao = str(data.get("observacao", "")).strip()
    posto_cfg = get_posto_config(posto_key)

    if not posto_cfg or not chassi or acao not in ["iniciar", "parar", "finalizar"]:
        return {"status": "erro", "detail": "Dados inválidos"}
    if get_posto_mode(posto_key) != "operacao":
        return {"status": "erro", "detail": "Este posto usa checklist de B.O.M."}
    if not can_access_posto(request, posto_key):
        return {"status": "erro", "detail": "Posto não permitido para este usuário"}

    veiculo = get_vehicle_by_chassi(db, chassi)
    if not veiculo:
        return {"status": "erro", "detail": "Chassi não encontrado"}

    apont = get_or_create_apontamento(db, chassi, posto_cfg["etapa"])
    status_atual = build_operacao_status(apont)
    if status_atual == "FINALIZADO":
        return {"status": "erro", "detail": "Este card já foi finalizado"}

    agora = datetime.datetime.now(LOCAL_TZ)
    responsavel = user.get("nome", "")

    if acao == "iniciar":
        if not apont.inicio:
            apont.inicio = agora
        apont.termino = None
        apont.status = "NÃO"
        registrar_historico_evento(db, veiculo, posto_cfg["etapa"], "NÃO", responsavel, inicio=apont.inicio, termino=apont.termino, observacao=observacao)
    elif acao == "parar":
        if not apont.inicio:
            apont.inicio = agora
        apont.termino = agora
        apont.status = "NÃO"
        registrar_historico_evento(db, veiculo, posto_cfg["etapa"], "NÃO", responsavel, inicio=apont.inicio, termino=apont.termino, observacao=observacao)
    elif acao == "finalizar":
        if not apont.inicio:
            apont.inicio = agora
        apont.termino = agora
        apont.status = "SIM"
        registrar_historico_evento(db, veiculo, posto_cfg["etapa"], "SIM", responsavel, inicio=apont.inicio, termino=apont.termino, observacao=observacao)

    apont.responsavel = responsavel
    apont.localizacao = veiculo.localizacao if veiculo else None
    apont.observacao = observacao
    db.commit()
    return {"status": "ok"}


@app.post("/bom/item-status")
async def bom_item_status(request: Request, data: dict = Body(...), db: Session = Depends(database.get_db)):
    user = require_login(request)
    if not user:
        return {"status": "erro", "detail": "Login necessário"}
    if is_management_user(request):
        return {"status": "erro", "detail": "Ação disponível apenas na visão operacional"}

    item_id = data.get("item_id")
    status = normalize_lookup_key(data.get("status"))
    if status not in ["SIM", "NAO", "N_A"]:
        return {"status": "erro", "detail": "Status inválido"}

    item = get_bom_item(db, item_id)
    if not item:
        return {"status": "erro", "detail": "Item não encontrado"}
    if not can_access_chassi(request, db, item.chassi):
        return {"status": "erro", "detail": "Card não permitido para este usuário"}

    item.status = normalize_status_value(status)
    item.responsavel = user.get("nome", "")
    item.atualizado_em = datetime.datetime.now(LOCAL_TZ)
    sync_stage_from_bom(db, item.tipo, item.chassi, user.get("nome", ""))
    db.commit()
    return {"status": "ok"}


@app.post("/bom/item-quantidade")
async def bom_item_quantidade(request: Request, data: dict = Body(...), db: Session = Depends(database.get_db)):
    user = require_login(request)
    if not user:
        return {"status": "erro", "detail": "Login necessario"}
    if is_management_user(request):
        return {"status": "erro", "detail": "Acao disponivel apenas na visao operacional"}
    if not user:
        return {"status": "erro", "detail": "Login necessÃ¡rio"}
    if is_management_user(request):
        return {"status": "erro", "detail": "AÃ§Ã£o disponÃ­vel apenas na visÃ£o operacional"}

    item = get_bom_item(db, data.get("item_id"))
    if not item:
        return {"status": "erro", "detail": "Item nao encontrado"}
    if not can_access_chassi(request, db, item.chassi):
        return {"status": "erro", "detail": "Card nao permitido para este usuario"}
    if not item:
        return {"status": "erro", "detail": "Item nÃ£o encontrado"}
    if not can_access_chassi(request, db, item.chassi):
        return {"status": "erro", "detail": "Card nÃ£o permitido para este usuÃ¡rio"}

    quantidade = parse_quantity_value(data.get("quantidade"))
    if quantidade is None or quantidade < 0:
        return {"status": "erro", "detail": "Informe uma quantidade valida para apontamento"}

    item.qtd_apontada = quantidade
    item.responsavel = user.get("nome", "")
    item.atualizado_em = datetime.datetime.now(LOCAL_TZ)
    db.commit()
    return {"status": "ok", "quantidade": format_quantity_value(quantidade)}


@app.post("/bom/empenho")
async def bom_empenho(request: Request, data: dict = Body(...), db: Session = Depends(database.get_db)):
    user = require_login(request)
    if not user:
        return {"status": "erro", "detail": "Login necessário"}
    if is_management_user(request):
        return {"status": "erro", "detail": "Ação disponível apenas na visão operacional"}

    item = get_bom_item(db, data.get("item_id"))
    if not item:
        return {"status": "erro", "detail": "Item não encontrado"}
    if str(item.tipo).upper() != "EXPEDICAO":
        return {"status": "erro", "detail": "Empenho disponível apenas para expedição"}
    if not can_access_chassi(request, db, item.chassi):
        return {"status": "erro", "detail": "Card não permitido para este usuário"}

    qtd_empenhada = parse_quantity_value(data.get("quantidade"))
    if qtd_empenhada is None or qtd_empenhada <= 0:
        return {"status": "erro", "detail": "Informe uma quantidade valida para o empenho"}

    total_lancado_atual = db.query(func.coalesce(func.sum(models.Empenho.qtd_empenhada), 0.0)).filter(
        models.Empenho.bom_item_id == item.id
    ).scalar() or 0.0
    item.qtd_apontada = resolve_bom_item_pointed_quantity(item, total_lancado_atual) + qtd_empenhada
    item.responsavel = user.get("nome", "")
    item.atualizado_em = datetime.datetime.now(LOCAL_TZ)

    db.add(
        models.Empenho(
            bom_item_id=item.id,
            chassi=item.chassi,
            cod_item=item.cod_item,
            item=item.item,
            descricao=item.descricao,
            qtd_empenhada=qtd_empenhada,
            sequencia_producao=get_sequence_number(db, "EXPEDICAO", item.chassi),
            responsavel=user.get("nome", ""),
        )
    )
    db.commit()
    return {
        "status": "ok",
        "quantidade": format_quantity_value(qtd_empenhada),
        "qtd_apontada": format_quantity_value(item.qtd_apontada),
    }


@app.get("/os/{chassi}", response_class=HTMLResponse)
async def ordem_servico_view(request: Request, chassi: str, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not can_access_chassi(request, db, chassi):
        return RedirectResponse(url=get_operator_home_url(request), status_code=303)

    ordem_servico = get_ordem_servico(db, chassi)
    if not ordem_servico:
        return RedirectResponse(url=get_operator_home_url(request), status_code=303)

    preview = extract_docx_preview(ordem_servico.caminho_arquivo)
    return templates.TemplateResponse(
        request,
        "os_preview.html",
        {
            "request": request,
            "current_user": require_login(request),
            "ordem_servico": ordem_servico,
            "preview": preview,
        },
    )

@app.get("/sequenciamento", response_class=HTMLResponse)
async def sequenciamento_page(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)
    return render_sequenciamento_page(request, db)

@app.post("/sequenciamento", response_class=HTMLResponse)
async def sequenciamento_save(
    request: Request,
    posto: str = Form(...),
    chassi: str = Form(...),
    sequencia: int = Form(...),
    db: Session = Depends(database.get_db),
):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)

    posto_key = str(posto).strip().upper()
    chassi_key = str(chassi).strip()
    form_data = {"posto": posto_key, "chassi": chassi_key, "sequencia": sequencia}

    if not get_posto_config(posto_key):
        return render_sequenciamento_page(request, db, erro="Selecione um posto válido.", form_data=form_data)
    if sequencia < 1:
        return render_sequenciamento_page(request, db, erro="A sequência deve ser maior que zero.", form_data=form_data)

    veiculo = get_vehicle_by_chassi(db, chassi_key)
    if not veiculo:
        return render_sequenciamento_page(request, db, erro="Chassi não encontrado na base atual.", form_data=form_data)

    resequence_posto(db, posto_key, chassi_key, sequencia)
    db.commit()
    return render_sequenciamento_page(request, db, sucesso="Sequenciamento salvo com sucesso.")


@app.post("/sequenciamento/automatico", response_class=HTMLResponse)
async def sequenciamento_automatico(
    request: Request,
    db: Session = Depends(database.get_db),
):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)

    generate_automatic_sequences(db)
    db.commit()
    return render_sequenciamento_page(
        request,
        db,
        sucesso="Sequenciamento automatico gerado com base na logica principal. Agora voce pode editar ou excluir ajustes pontuais.",
    )


@app.post("/sequenciamento/excluir", response_class=HTMLResponse)
async def sequenciamento_delete(
    request: Request,
    posto: str = Form(...),
    chassi: str = Form(...),
    db: Session = Depends(database.get_db),
):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)

    posto_key = str(posto).strip().upper()
    chassi_key = str(chassi).strip()
    form_data = {"posto": posto_key, "chassi": chassi_key}

    if not posto_key or not chassi_key:
        return render_sequenciamento_page(request, db, erro="Informe o posto e o chassi para excluir a atribuicao.", form_data=form_data)

    removed = remove_posto_sequencia(db, posto_key, chassi_key)
    if not removed:
        return render_sequenciamento_page(request, db, erro="Atribuicao nao encontrada para exclusao.", form_data=form_data)

    db.commit()
    return render_sequenciamento_page(request, db, sucesso="Atribuicao removida com sucesso.")


@app.post("/sequenciamento/os-upload", response_class=HTMLResponse)
async def sequenciamento_os_upload(
    request: Request,
    chassi: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(database.get_db),
):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)

    chassi_key = str(chassi).strip()
    if not chassi_key:
        return render_sequenciamento_page(request, db, erro="Informe o chassi da O.S.")
    if not file.filename or not file.filename.lower().endswith(".docx"):
        return render_sequenciamento_page(request, db, erro="A O.S. deve ser enviada em arquivo DOCX.")

    veiculo = get_vehicle_by_chassi(db, chassi_key)
    if not veiculo:
        return render_sequenciamento_page(request, db, erro="Chassi não encontrado na base atual.")

    content = await file.read()
    if not content:
        return render_sequenciamento_page(request, db, erro="Arquivo vazio.")

    existing = get_ordem_servico(db, chassi_key)
    if existing:
        remove_ordem_servico_arquivo(existing)

    safe_name = f"{chassi_key}_{int(datetime.datetime.now(LOCAL_TZ).timestamp())}.docx"
    destino = OS_UPLOAD_DIR / safe_name
    destino.write_bytes(content)

    if existing:
        existing.nome_arquivo = file.filename
        existing.caminho_arquivo = str(destino)
    else:
        db.add(
            models.OrdemServico(
                chassi=chassi_key,
                nome_arquivo=file.filename,
                caminho_arquivo=str(destino),
            )
        )

    composition_rows = extract_docx_composition_items(str(destino))
    if composition_rows:
        prep_count = sync_bom_items_for_chassi(db, "PREPARACAO", chassi_key, composition_rows)
        exp_count = sync_bom_items_for_chassi(db, "EXPEDICAO", chassi_key, composition_rows)
        sync_stage_from_bom(db, "PREPARACAO", chassi_key, "Sistema")
        sync_stage_from_bom(db, "EXPEDICAO", chassi_key, "Sistema")
        db.commit()
        return render_sequenciamento_page(
            request,
            db,
            sucesso=f"O.S. enviada com sucesso. Checklist atualizado com {prep_count} item(ns) para Preparacao e {exp_count} para Expedicao.",
        )

    db.commit()
    return render_sequenciamento_page(
        request,
        db,
        sucesso="O.S. enviada com sucesso. Nenhuma tabela de composicao valida foi encontrada para gerar o checklist.",
    )


@app.post("/sequenciamento/bom-upload", response_class=HTMLResponse)
async def sequenciamento_bom_upload(
    request: Request,
    tipo: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(database.get_db),
):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)

    tipo_key = str(tipo).strip().upper()
    if tipo_key not in BOM_TIPOS:
        return render_sequenciamento_page(request, db, erro="Tipo de B.O.M. inválido.")

    content = await file.read()
    try:
        df = (
            pd.read_excel(io.BytesIO(content))
            if file.filename and file.filename.lower().endswith(".xlsx")
            else pd.read_csv(io.BytesIO(content))
        )
    except Exception as exc:
        return render_sequenciamento_page(request, db, erro=f"Erro ao ler B.O.M.: {exc}")

    df.columns = [normalize_lookup_key(c) for c in df.columns]
    aliases = {
        "CHASSI": ["N_CHASSI", "NO_CHASSI", "NUMERO_CHASSI", "CHASSI"],
        "COD_ITEM": ["COD_ITEM"],
        "ITEM": ["ITEM"],
        "DESCRICAO": ["DESCRICAO", "DESCRI_O", "DESCRI__O"],
        "QTD": ["QTD", "QUANTIDADE"],
    }

    def pick(row, key):
        for col in aliases[key]:
            if col in df.columns:
                return safe_str(row.get(col, ""))
        return ""

    rows = []
    chassis_in_upload = set()
    for _, row in df.iterrows():
        chassi_row = pick(row, "CHASSI").split(".")[0]
        if not chassi_row:
            continue
        rows.append(
            {
                "chassi": chassi_row,
                "cod_item": pick(row, "COD_ITEM"),
                "item": pick(row, "ITEM"),
                "descricao": pick(row, "DESCRICAO"),
                "qtd": pick(row, "QTD"),
            }
        )
        chassis_in_upload.add(chassi_row)

    if not rows:
        return render_sequenciamento_page(request, db, erro="Nenhuma linha válida encontrada na B.O.M.")

    itens_antigos = db.query(models.BomItem).filter(
        models.BomItem.tipo == tipo_key,
        models.BomItem.chassi.in_(list(chassis_in_upload)),
    ).all()
    itens_antigos_ids = [item.id for item in itens_antigos]
    if itens_antigos_ids:
        db.query(models.Empenho).filter(
            models.Empenho.bom_item_id.in_(itens_antigos_ids)
        ).delete(synchronize_session=False)
    db.query(models.BomItem).filter(
        models.BomItem.tipo == tipo_key,
        models.BomItem.chassi.in_(list(chassis_in_upload)),
    ).delete(synchronize_session=False)

    posto_key, posto_cfg = get_posto_by_bom_tipo(tipo_key)
    for row in rows:
        db.add(
            models.BomItem(
                tipo=tipo_key,
                chassi=row["chassi"],
                cod_item=row["cod_item"],
                item=row["item"],
                descricao=row["descricao"],
                qtd=row["qtd"],
                status="NAO",
            )
        )
        apont = get_or_create_apontamento(db, row["chassi"], posto_cfg["etapa"])
        apont.status = "NAO"
        apont.termino = None

    db.commit()
    return render_sequenciamento_page(request, db, sucesso=f"{BOM_TIPOS[tipo_key]} enviada com sucesso.")

@app.get("/importar")
async def pg_importar(request: Request):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)
    return templates.TemplateResponse(
        request,
        "importar.html",
        {
            "request": request,
            "current_user": require_login(request),
            "is_admin": bool(require_admin(request)),
            "localizacoes": LOCALIZACOES,
        },
    )

@app.get("/usuarios", response_class=HTMLResponse)
async def usuarios_page(request: Request, db: Session = Depends(database.get_db)):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)
    return render_user_management(request, db)

@app.post("/usuarios", response_class=HTMLResponse)
async def usuarios_create(
    request: Request,
    nome: str = Form(...),
    login: str = Form(...),
    senha: str = Form(...),
    perfil: str = Form(...),
    db: Session = Depends(database.get_db),
):
    if not require_login(request):
        return RedirectResponse(url="/login", status_code=303)
    if not require_admin(request):
        return RedirectResponse(url="/", status_code=303)

    nome_limpo = str(nome).strip()
    login_limpo = normalize_login(login)
    senha_limpa = str(senha).strip()
    perfil_limpo = normalize_profile(perfil)
    form_data = {"nome": nome_limpo, "login": login_limpo, "perfil": perfil_limpo}

    if not nome_limpo or not login_limpo or not senha_limpa or not perfil_limpo:
        return render_user_management(
            request,
            db,
            erro="Preencha nome, login, senha e perfil.",
            form_data=form_data,
        )

    existente = db.query(models.Usuario).filter(func.lower(models.Usuario.login) == login_limpo).first()
    if existente:
        return render_user_management(
            request,
            db,
            erro="Ja existe um usuario com esse login.",
            form_data=form_data,
        )

    db.add(
        models.Usuario(
            nome=nome_limpo,
            login=login_limpo,
            senha_hash=hash_password(senha_limpa),
            perfil=perfil_limpo,
        )
    )
    db.commit()
    return render_user_management(request, db, sucesso="Usuario criado com sucesso.")

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request, db: Session = Depends(database.get_db)):
    if require_login(request):
        return RedirectResponse(url=get_operator_home_url(request), status_code=303)
    return render_login_page(request, db)

@app.post("/login")
async def login_post(
    request: Request,
    login: str = Form(...),
    senha: str = Form(...),
    db: Session = Depends(database.get_db),
):
    login_limpo = normalize_login(login)
    senha_limpa = str(senha).strip()
    if not login_limpo or not senha_limpa:
        return render_login_page(request, db, erro="Informe login e senha.", login_value=login_limpo)

    user = db.query(models.Usuario).filter(func.lower(models.Usuario.login) == login_limpo).first()
    if not user or not verify_password(senha_limpa, user.senha_hash):
        return render_login_page(request, db, erro="Login ou senha invalidos.", login_value=login_limpo)

    request.session.clear()
    request.session["user"] = build_session_user(user)
    return RedirectResponse(url=get_operator_home_url_for_profile(user.perfil), status_code=303)

@app.get("/logout")
async def logout(request: Request):
    request.session.clear()
    return RedirectResponse(url="/login", status_code=303)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8001))
    host = os.environ.get("HOST", "127.0.0.1")
    uvicorn.run(app, host=host, port=port)
